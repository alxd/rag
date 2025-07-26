import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import pandas as pd
from sklearn.preprocessing import MultiLabelBinarizer
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from upsetplot import UpSet, from_indicators
import numpy as np
import os
import warnings
from PIL import Image, ImageTk, ImageDraw, ImageFont
import threading
from collections import defaultdict
import re
import html
import traceback
import textwrap
import win32com.client
import subprocess
import random
import json
import openai
try:
    from mistralai import Mistral
except ImportError:
    Mistral = None
try:
    from huggingface_hub import InferenceClient
except ImportError:
    InferenceClient = None

# Common suffixes for substring grouping (except suffixes)
common_suffixes = [
    'ation','ption', 'ment', 'ness', 'sion', 'tion', 'ing', 'ed', 'ly', 'er', 'est', 'ful', 'less', 'able', 'ible', 'ous', 'ive', 'al', 'ic', 'ant', 'ent', 'ism', 'ist', 'ity', 'ty', 'en', 'ize', 'ise', 'ward', 'wise'
]

warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

def wrap_label(text, width=18):
    # Try to wrap at word boundaries, fallback to hard wrap
    return '\n'.join(textwrap.wrap(text, width=width))

class ZoomablePanCanvas:
    def __init__(self, parent):
        self.parent = parent
        self.canvas = tk.Canvas(parent, bg='white')
        
        # Scrollbars
        self.h_scrollbar = ttk.Scrollbar(parent, orient="horizontal", command=self.canvas.xview)
        self.v_scrollbar = ttk.Scrollbar(parent, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(xscrollcommand=self.h_scrollbar.set, yscrollcommand=self.v_scrollbar.set)
        
        # Grid layout
        self.canvas.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.h_scrollbar.grid(row=1, column=0, sticky=(tk.W, tk.E))
        self.v_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # Configure grid weights
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)
        
        # Zoom and pan variables
        self.zoom_level = 1.0
        self.original_images = []
        self.image_refs = []
        self.image_items = []
        
        # Bind mouse events for panning
        self.canvas.bind("<Button-1>", self.start_pan)
        self.canvas.bind("<B1-Motion>", self.do_pan)
        self.canvas.bind("<MouseWheel>", self.zoom_wheel)
        
        self.last_x = 0
        self.last_y = 0
        
    def start_pan(self, event):
        self.canvas.scan_mark(event.x, event.y)
        self.last_x = event.x
        self.last_y = event.y
        
    def do_pan(self, event):
        self.canvas.scan_dragto(event.x, event.y, gain=1)
        
    def zoom_wheel(self, event):
        if event.delta > 0:
            self.zoom_in()
        else:
            self.zoom_out()
            
    def set_images(self, image_paths, window_width):
        self.original_images = []
        self.image_refs = []
        self.image_items = []
        self.canvas.delete("all")
        if not image_paths:
            return
        loaded_imgs = []
        max_height = 0
        for img_path in image_paths:
            try:
                img = Image.open(img_path)
                loaded_imgs.append(img)
                max_height = max(max_height, img.height)
            except Exception as e:
                print(f"Error loading image {img_path}: {e}")
        x_offset = 10
        padding = 10
        for img in loaded_imgs:
            # Align to bottom
            y_offset = max_height - img.height + padding
            self.original_images.append((img, x_offset, y_offset))
            x_offset += img.width + padding
        self.zoom_level = 1.0
        self.update_display()
        
    def update_display(self):
        self.canvas.delete("all")
        self.image_refs = []
        self.image_items = []
        
        for img, x, y in self.original_images:
            # Apply zoom
            new_width = int(img.width * self.zoom_level)
            new_height = int(img.height * self.zoom_level)
            
            if new_width > 0 and new_height > 0:
                zoomed_img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(zoomed_img)
                self.image_refs.append(photo)
                
                # Position with zoom
                scaled_x = int(x * self.zoom_level)
                scaled_y = int(y * self.zoom_level)
                
                item = self.canvas.create_image(scaled_x, scaled_y, anchor="nw", image=photo)
                self.image_items.append(item)
        
        # Update scroll region
        self.canvas.update_idletasks()
        bbox = self.canvas.bbox("all")
        if bbox:
            self.canvas.configure(scrollregion=bbox)
            
    def zoom_in(self):
        self.zoom_level = min(self.zoom_level * 1.1, 5.0)
        self.update_display()
        
    def zoom_out(self):
        self.zoom_level = max(self.zoom_level / 1.1, 0.2)
        self.update_display()
        
    def fit_to_window(self):
        if not self.original_images:
            return
        # Calculate total width of all images + padding
        total_width = sum(img.width for img, _, _ in self.original_images)
        total_width += 10 * (len(self.original_images) + 1)
        canvas_width = self.canvas.winfo_width()
        if canvas_width <= 1:  # Not yet rendered
            self.parent.update_idletasks()
            canvas_width = self.canvas.winfo_width()
        if total_width > 0 and canvas_width > 0:
            self.zoom_level = min(1.0, canvas_width / total_width)
        else:
            self.zoom_level = 1.0
        self.update_display()

class UpSetGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("UpSet Plot Generator")
        self.root.geometry("1400x800")
        
        self.csv_file = None
        self.output_images = []
        self.color_mapping = {}
        self.use_colors = tk.BooleanVar(value=True)
        self.group_by_subletters = tk.BooleanVar(value=False)
        self.group_by_words = tk.BooleanVar(value=True)
        self.group_by_same_color = tk.BooleanVar(value=False)
        
        # New variables for refined color grouping
        self.agg_use_colors = tk.BooleanVar(value=True)
        self.agg_group_by_subletters = tk.BooleanVar(value=False)
        self.agg_group_by_words = tk.BooleanVar(value=True)
        self.agg_enable_fuzzy = tk.BooleanVar(value=True)
        
        self.setup_ui()
        
    def setup_ui(self):
        # Main frame
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.columnconfigure(2, weight=1)
        main_frame.rowconfigure(5, weight=1)  # Make results frame expand
        
        # --- Merge CSVs section ---
        merge_frame = ttk.LabelFrame(main_frame, text="Merge CSV Files", padding="5")
        merge_frame.grid(row=0, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0,2))
        self.merge_label = ttk.Label(merge_frame, text="No files selected", foreground="gray")
        self.merge_label.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5), pady=2)
        # Place the button below the label, natural width
        self.merge_button = ttk.Button(merge_frame, text="Select & Merge CSVs", command=self.merge_csv_files)
        self.merge_button.grid(row=1, column=0, sticky=tk.W, padx=(0, 5), pady=2)
        
        # File selection
        ttk.Label(main_frame, text="CSV File:").grid(row=1, column=0, sticky=tk.W, pady=0)
        self.file_label = ttk.Label(main_frame, text="No file selected", foreground="gray")
        self.file_label.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=(5, 0), pady=0)
        ttk.Button(main_frame, text="Browse", command=self.browse_file).grid(row=1, column=2, padx=(5, 0), pady=0)
        # Block info label (single line, under file label)
        self.block_info_label = ttk.Label(main_frame, text="", foreground="blue")
        self.block_info_label.grid(row=2, column=1, sticky=(tk.W, tk.E), padx=(5, 0), pady=0)
        
        # Options frame
        options_frame = ttk.LabelFrame(main_frame, text="Options", padding="3")
        options_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=0)
        
        # Color checkbox and entry
        color_frame = ttk.Frame(options_frame)
        color_frame.grid(row=0, column=0, sticky=tk.W, pady=0)
        self.color_criteria_cb = ttk.Checkbutton(color_frame, text="Color criteria", variable=self.use_colors, command=self.on_color_criteria_toggle)
        self.color_criteria_cb.grid(row=0, column=0, sticky=tk.W, pady=0)
        ttk.Label(color_frame, text="Min shared letters:").grid(row=0, column=1, padx=(5,0), pady=0)
        self.min_letters_var = tk.StringVar(value="5")
        self.min_letters_entry = ttk.Entry(color_frame, textvariable=self.min_letters_var, width=3)
        self.min_letters_entry.grid(row=0, column=2, padx=(2,0), pady=0)
        ttk.Checkbutton(color_frame, text="Group by subletters (except suffixes)", variable=self.group_by_subletters, command=self.on_group_by_subletters).grid(row=0, column=3, padx=(5,0), pady=0)
        ttk.Checkbutton(color_frame, text="Group by whole words", variable=self.group_by_words, command=self.on_group_by_words).grid(row=0, column=4, padx=(5,0), pady=0)
        self.group_by_same_color_cb = ttk.Checkbutton(color_frame, text="Collapse into same color group", variable=self.group_by_same_color)
        self.group_by_same_color_cb.grid(row=0, column=5, padx=(5,0), pady=0)
        self.group_by_same_color_cb.state(['disabled'])
        # Add label wrap width entry
        ttk.Label(color_frame, text="Label wrap width:").grid(row=0, column=6, padx=(5,0), pady=0)
        self.wrap_width_var = tk.StringVar(value="75")
        self.wrap_width_entry = ttk.Entry(color_frame, textvariable=self.wrap_width_var, width=3)
        self.wrap_width_entry.grid(row=0, column=7, padx=(2,0), pady=0)
        
        # Process button and progress
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=0)
        button_frame.columnconfigure(0, weight=1)
        
        self.progress = ttk.Progressbar(button_frame, mode='indeterminate')
        self.progress.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5), pady=0)
        
        self.process_btn = ttk.Button(button_frame, text="Generate UpSet Plots", 
                                     command=self.start_processing, state='disabled')
        self.process_btn.grid(row=0, column=1, pady=0)
        
        # --- Aggregate Results section ---
        aggregate_frame = ttk.LabelFrame(main_frame, text="Aggregate Results from Multiple Folders", padding="5")
        aggregate_frame.grid(row=7, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(10,2))
        self.aggregate_folder = None
        self.aggregate_folder_label = ttk.Label(aggregate_frame, text="No folder selected", foreground="gray")
        self.aggregate_folder_label.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5), pady=2)
        self.select_aggregate_folder_btn = ttk.Button(aggregate_frame, text="Select Parent Folder", command=self.select_aggregate_folder)
        self.select_aggregate_folder_btn.grid(row=1, column=0, sticky=tk.W, padx=(0, 5), pady=2)
        self.aggregate_btn = ttk.Button(aggregate_frame, text="Aggregate Results", command=self.aggregate_results, state='disabled')
        self.aggregate_btn.grid(row=1, column=1, sticky=tk.W, padx=(5, 0), pady=2)
        self.aggregate_status_label = ttk.Label(aggregate_frame, text="", foreground="blue")
        self.aggregate_status_label.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=(0, 5), pady=2)
        # --- Progress bar for aggregation ---
        self.aggregate_progress = ttk.Progressbar(aggregate_frame, mode='determinate', length=300)
        self.aggregate_progress.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(2,2))
        self.aggregate_time_label = ttk.Label(aggregate_frame, text="", foreground="gray")
        self.aggregate_time_label.grid(row=5, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=(0, 5), pady=2)
        
        # --- LLM Grouping Controls ---
        llm_grouping_frame = ttk.Frame(aggregate_frame)
        llm_grouping_frame.grid(row=6, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(5,2))
        self.llm_grouping_var = tk.BooleanVar(value=False)
        self.llm_grouping_cb = ttk.Checkbutton(llm_grouping_frame, text="LLM Grouping (semantic)", variable=self.llm_grouping_var, command=self.on_llm_grouping_toggle)
        self.llm_grouping_cb.grid(row=0, column=0, sticky=tk.W, pady=0)
        self.llm_model_var = tk.StringVar(value="Mistral-API")
        self.llm_model_dropdown = ttk.Combobox(llm_grouping_frame, textvariable=self.llm_model_var, values=[
            "Remote Meta-Llama-3", "Mistral-API", "GPT-3.5", "GPT-4o", "GPT-4o mini", "o1-mini", "o3-mini",
            "Gemini", "Claude", "Grok", "Qwen3", "Phi4", "Meta Llama 70B", "DeepSeek V3", "Mistral (Nebius)"
        ], state="readonly", width=18)
        self.llm_model_dropdown.grid(row=0, column=1, padx=(10,0), pady=0)
        self.llm_prompt_var = tk.StringVar(value="Group the following concepts by meaning. Return the result as JSON, where each group is a list of concepts.\nConcepts:\n- concept1\n- concept2\n...")
        self.llm_prompt_entry = ttk.Entry(llm_grouping_frame, textvariable=self.llm_prompt_var, width=60)
        self.llm_prompt_entry.grid(row=0, column=2, padx=(10,0), pady=0)
        # --- Refined Color Grouping Controls ---
        refined_grouping_frame = ttk.Frame(aggregate_frame)
        refined_grouping_frame.grid(row=7, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(5,2))
        
        # Color grouping options
        agg_color_frame = ttk.Frame(refined_grouping_frame)
        agg_color_frame.grid(row=0, column=0, sticky=tk.W, pady=2)
        self.agg_color_criteria_cb = ttk.Checkbutton(agg_color_frame, text="Refined Color Grouping", variable=self.agg_use_colors, command=self.on_agg_color_criteria_toggle)
        self.agg_color_criteria_cb.grid(row=0, column=0, sticky=tk.W, pady=0)
        ttk.Label(agg_color_frame, text="Min shared letters:").grid(row=0, column=1, padx=(5,0), pady=0)
        self.agg_min_letters_var = tk.StringVar(value="5")
        self.agg_min_letters_entry = ttk.Entry(agg_color_frame, textvariable=self.agg_min_letters_var, width=3)
        self.agg_min_letters_entry.grid(row=0, column=2, padx=(2,0), pady=0)
        ttk.Checkbutton(agg_color_frame, text="Group by subletters (except suffixes)", variable=self.agg_group_by_subletters, command=self.on_agg_group_by_subletters).grid(row=0, column=3, padx=(5,0), pady=0)
        ttk.Checkbutton(agg_color_frame, text="Group by whole words", variable=self.agg_group_by_words, command=self.on_agg_group_by_words).grid(row=0, column=4, padx=(5,0), pady=0)
        
        # Fuzzy logic controls
        fuzzy_frame = ttk.Frame(refined_grouping_frame)
        fuzzy_frame.grid(row=1, column=0, sticky=tk.W, pady=2)
        ttk.Checkbutton(fuzzy_frame, text="Enable Fuzzy Logic", variable=self.agg_enable_fuzzy, command=self.on_agg_enable_fuzzy_toggle).grid(row=0, column=0, sticky=tk.W, pady=0)
        ttk.Label(fuzzy_frame, text="Similarity Threshold:").grid(row=0, column=1, padx=(5,0), pady=0)
        self.sim_threshold_var = tk.DoubleVar(value=0.85)
        self.sim_threshold_entry = ttk.Entry(fuzzy_frame, textvariable=self.sim_threshold_var, width=5)
        self.sim_threshold_entry.grid(row=0, column=2, sticky=tk.W, padx=(2,10))
        ttk.Label(fuzzy_frame, text="Grouping Logic:").grid(row=0, column=3, sticky=tk.W)
        self.grouping_logic_var = tk.StringVar(value="Fuzzy")
        self.grouping_logic_combo = ttk.Combobox(fuzzy_frame, textvariable=self.grouping_logic_var, values=["Fuzzy", "Exact", "None"], state="readonly", width=8)
        self.grouping_logic_combo.grid(row=0, column=4, sticky=tk.W, padx=(2,0))
        # Add fuzzy logic explanation label
        self.fuzzy_explanation_label = ttk.Label(fuzzy_frame, text="Fuzzy logic groups concepts by overall string similarity.\nThreshold 1.0 = only identical concepts.\nThreshold 0.85 = minor spelling/word order differences.\nThreshold 0.6 = allows more distant matches.\nExample: 'justice' and 'justices' are grouped at 0.85, but 'justice' and 'injustice' only at 0.6.", justify='left', foreground='gray')
        self.fuzzy_explanation_label.grid(row=1, column=0, columnspan=5, sticky=tk.W, pady=(2,0))
        
        # Initialize fuzzy controls state
        self.on_agg_enable_fuzzy_toggle()
        
        # Results frame with zoom controls
        results_frame = ttk.LabelFrame(main_frame, text="Results", padding="5")
        results_frame.grid(row=5, column=0, columnspan=3, sticky=(tk.N, tk.S, tk.E, tk.W), pady=5)
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(1, weight=1)
        
        # Zoom controls
        zoom_frame = ttk.Frame(results_frame)
        zoom_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 2))
        
        ttk.Button(zoom_frame, text="Zoom In", command=self.zoom_in).grid(row=0, column=0, padx=(0, 2), pady=0)
        ttk.Button(zoom_frame, text="Zoom Out", command=self.zoom_out).grid(row=0, column=1, padx=(0, 2), pady=0)
        ttk.Button(zoom_frame, text="Fit to Window", command=self.fit_to_window).grid(row=0, column=2, padx=(0, 2), pady=0)
        
        ttk.Label(zoom_frame, text="Use mouse wheel to zoom, click and drag to pan").grid(row=0, column=3, padx=(10, 0), pady=0)
        
        # Zoomable canvas
        canvas_frame = ttk.Frame(results_frame)
        canvas_frame.grid(row=1, column=0, sticky=(tk.N, tk.S, tk.E, tk.W))
        canvas_frame.columnconfigure(0, weight=1)
        canvas_frame.rowconfigure(0, weight=1)
        
        self.zoomable_canvas = ZoomablePanCanvas(canvas_frame)
        self.zoomable_canvas.canvas.grid(sticky=(tk.N, tk.S, tk.E, tk.W))  # Ensure canvas expands
        
        # Status label
        self.status_label = ttk.Label(main_frame, text="Ready")
        self.status_label.grid(row=6, column=0, columnspan=3, pady=2)
        
    def browse_file(self):
        file_path = filedialog.askopenfilename(
            title="Select CSV file",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        
        if file_path:
            self.csv_file = file_path
            self.file_label.config(text=os.path.basename(file_path), foreground="black")
            self.process_btn.config(state='normal')
            # Try to read and display block info
            self.display_block_info()
            
    def display_block_info(self):
        # Try to read the CSV and infer block structure
        if not self.csv_file:
            self.block_info_label.config(text="", foreground="blue")
            return
        encodings_to_try = ['cp1253', 'utf-8', 'windows-1252']
        df = None
        for enc in encodings_to_try:
            try:
                df = pd.read_csv(self.csv_file, encoding=enc)
                break
            except Exception:
                continue
        if df is None:
            self.block_info_label.config(text="Could not read CSV to infer blocks.", foreground="red")
            return
        # Detect block structure
        param_names = [col for col in ['Temperature', 'Top-p', 'Top-k', 'BM25 Weight'] if col in df.columns]
        block_size = None
        blocks = []
        if len(param_names) == 4:
            n = len(df)
            for p in param_names:
                values = df[p].values
                first_val = values[0]
                block_len = 1
                for v in values[1:]:
                    if v == first_val:
                        block_len += 1
                    else:
                        break
                if n % block_len == 0 and block_len > 1:
                    block_size = block_len
                    break
            if block_size is None:
                block_info = "Could not infer block size."
            else:
                num_blocks = n // block_size
                blocks = []
                block_strs = []
                for i in range(num_blocks):
                    start = i * block_size
                    end = (i + 1) * block_size
                    varying_param = None
                    for p in param_names:
                        if len(df[p].iloc[start:end].unique()) > 1:
                            varying_param = p
                            break
                    blocks.append((start, end, varying_param if varying_param else '?'))
                    block_strs.append(f"{start}-{end-1}: {varying_param if varying_param else '?'}")
                block_info = "Blocks: [" + ", ".join(block_strs) + "]"
        else:
            block_info = "Could not detect all required parameters."
        self.block_info_label.config(text=block_info, foreground="blue")
        self.detected_blocks = blocks
            
    def zoom_in(self):
        self.zoomable_canvas.zoom_in()
        
    def zoom_out(self):
        self.zoomable_canvas.zoom_out()
        
    def fit_to_window(self):
        self.zoomable_canvas.fit_to_window()
        
    def extract_concepts(self, text):
        if pd.isna(text):
            return []
        lines = text.split('\n')
        concepts = []
        for line in lines:
            line = line.strip()
            if line.startswith(tuple(str(i)+'.' for i in range(1, 21))):
                # Markdown or plain numbered list
                if '**' in line:
                    c = line.split('**')[1].replace(':', '').replace('.', '').strip()
                    if c:
                        concepts.append(c)
                else:
                    c = line.lstrip('0123456789. ').replace(':', '').replace('.', '').strip()
                    if c:
                        concepts.append(c)
        return concepts
    
    def normalize_word(self, word):
        """Normalize word by removing plurals and common variations"""
        word = word.lower().strip()
        
        # Handle common plural forms
        if word.endswith('ies') and len(word) > 4:
            return word[:-3] + 'y'
        elif word.endswith('es') and len(word) > 3:
            return word[:-2]
        elif word.endswith('s') and len(word) > 2:
            return word[:-1]
        
        # Handle common suffixes
        suffixes = ['ing', 'ed', 'er', 'est', 'ly', 'tion', 'sion', 'ness', 'ment', 'able', 'ible', 'ful', 'less']
        for suffix in suffixes:
            if word.endswith(suffix) and len(word) > len(suffix) + 2:
                return word[:-len(suffix)]
        
        return word
    
    def extract_key_words(self, concept):
        """Extract meaningful words from a concept"""
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'should', 'could', 'can', 'may', 'might', 'must'}
        words = re.findall(r'\b\w+\b', concept.lower())
        meaningful_words = [self.normalize_word(word) for word in words if word not in stop_words and len(word) > 2]
        return meaningful_words
    
    def create_color_mapping(self, all_concepts):
        self.color_mapping = {}
        min_letters = 5
        try:
            min_letters = int(self.min_letters_var.get())
        except Exception:
            pass
        all_concepts_sorted = sorted(all_concepts, key=lambda x: x.lower())
        groups = []
        assigned = set()
        group_by_words = self.group_by_words.get()
        for i, concept in enumerate(all_concepts_sorted):
            concept_lc = concept.lower()
            if concept_lc in assigned:
                continue
            group = [concept]
            words1 = set(re.findall(r'\b\w+\b', concept_lc))
            for j, other in enumerate(all_concepts_sorted):
                other_lc = other.lower()
                if i == j or other_lc in assigned:
                    continue
                words2 = set(re.findall(r'\b\w+\b', other_lc))
                found = False
                if group_by_words:
                    if words1 & words2:
                        found = True
                else:
                    for w1 in words1:
                        for w2 in words2:
                            for k in range(len(w1) - min_letters + 1):
                                sub = w1[k:k+min_letters]
                                if sub and sub in w2:
                                    if (w1.endswith(sub) or w2.endswith(sub)) and sub in common_suffixes:
                                        continue
                                    found = True
                                    break
                            if found:
                                break
                        if found:
                            break
                if found:
                    group.append(other)
            groups.append(group)
            assigned.update([c.lower() for c in group])
        colors = [
            '#800000', '#FF8C00', '#228B22', '#8B008B', '#A0522D', '#2E8B57', '#9932CC', '#FFD700',
            '#556B2F', '#C71585', '#8B4513', '#20B2AA', '#B22222', '#FF4500', '#6A5ACD', '#D2691E',
            '#006400', '#708090', '#FF6347', '#483D8B', '#000000', '#808000', '#8B0000', '#FF1493',
        ]
        color_idx = 0
        for group in groups:
            color = colors[color_idx % len(colors)]
            for concept in group:
                self.color_mapping[concept] = color
            color_idx += 1
        return self.color_mapping
    
    def create_colored_text_image(self, text, color, font_size=12):
        """Create an image with colored text on white background"""
        try:
            font = ImageFont.truetype("arial.ttf", font_size)
        except:
            font = ImageFont.load_default()
        
        # Get text dimensions
        bbox = font.getbbox(text)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        # Create image with white background
        img = Image.new('RGB', (text_width + 10, text_height + 6), 'white')
        draw = ImageDraw.Draw(img)
        
        # Draw text in color
        draw.text((5, 3), text, fill=color, font=font)
        
        return img
    
    def start_processing(self):
        if not self.csv_file:
            messagebox.showerror("Error", "Please select a CSV file first")
            return
            
        # Start processing in a separate thread
        self.process_btn.config(state='disabled')
        self.progress.start()
        self.status_label.config(text="Processing...")
        
        thread = threading.Thread(target=self.process_data)
        thread.daemon = True
        thread.start()
        
    def process_data(self):
        try:
            print('Starting process_data')
            encodings_to_try = ['utf-8', 'windows-1252', 'cp1253', 'latin1']
            for enc in encodings_to_try:
                try:
                    df = pd.read_csv(self.csv_file, encoding=enc)
                    print(f'CSV loaded with encoding: {enc}')
                    break
                except Exception as e:
                    print(f'Failed to load CSV with encoding {enc}: {e}')
            else:
                raise Exception('Could not read CSV file with any supported encoding (utf-8, cp1253, windows-1252, latin1)')
            print('First 10 Main Answer values:')
            print(df['Main Answer'].head(10).to_list())
            df['Concepts'] = df['Main Answer'].apply(self.extract_concepts)
            print('First 10 Concepts values:')
            print(df['Concepts'].head(10).to_list())
            all_concepts = set()
            for concepts in df['Concepts']:
                all_concepts.update(concepts)
            print(f'All unique concepts: {all_concepts}')
            self.create_color_mapping(list(all_concepts))
            print(f'Color mapping: {self.color_mapping}')
            params = ['Temperature', 'Top-p', 'Top-k', 'BM25 Weight']
            param_labels = {'Temperature': 'Temp', 'Top-p': 'Topp', 'Top-k': 'Topk', 'BM25 Weight': 'BM25'}
            outdir = 'compare_gui_output'
            os.makedirs(outdir, exist_ok=True)
            blocks = getattr(self, 'detected_blocks', None)
            if not blocks or len(blocks) == 0:
                n = len(df)
                block_size = n // 4 if n % 4 == 0 else 5
                blocks = [(i*block_size, (i+1)*block_size, params[i] if i < len(params) else '?') for i in range(4)]
            output_files = []
            for start, end, varying_param in blocks:
                print(f'Processing block: {varying_param} ({start}-{end})')
                subset = df.iloc[start:end].copy()
                if subset.empty:
                    print('Subset empty, skipping')
                    continue
                mlb = MultiLabelBinarizer()
                concept_matrix = pd.DataFrame(mlb.fit_transform(subset['Concepts']), columns=mlb.classes_)
                print(f'Concept matrix columns: {concept_matrix.columns}')
                for p in params:
                    concept_matrix[p] = subset[p].values
                concept_matrix_reset = concept_matrix.drop(params, axis=1).astype(bool).reset_index(drop=True)
                print(f'Concept matrix reset columns: {concept_matrix_reset.columns}')
                if self.use_colors.get() and self.group_by_same_color.get():
                    color_map = {col: self.color_mapping.get(col, None) for col in concept_matrix_reset.columns}
                    print(f'Color map: {color_map}')
                    color_groups = defaultdict(list)
                    for col, color in color_map.items():
                        color_groups[color].append(col)
                    print(f'Color groups: {color_groups}')
                    merged = pd.DataFrame(index=concept_matrix_reset.index)
                    for color, cols in color_groups.items():
                        print(f'Grouping for color {color}: {cols}')
                        if color is None or len(cols) == 0:
                            continue
                        if len(cols) == 1:
                            merged[cols[0]] = concept_matrix_reset[cols[0]]
                        else:
                            group_name = "/".join(cols)
                            print(f'Group name: {group_name}')
                            group_name_wrapped = wrap_label(group_name, width=self.get_wrap_width())
                            merged[group_name_wrapped] = concept_matrix_reset[cols].any(axis=1)
                    concept_matrix_reset = merged
                print(f'Final columns for upset: {concept_matrix_reset.columns}')
                upset_data = from_indicators(concept_matrix_reset, concept_matrix_reset.columns)
                print('UpSet data created')
                fig = plt.figure(figsize=(12, 8))
                upset = UpSet(upset_data, show_counts=True)
                axes = upset.plot(fig=fig)
                bar_ax = axes['intersections']
                matrix_ax = axes['matrix']
                bars = bar_ax.patches
                upset_index = upset_data.index
                wrap_width = self.get_wrap_width()
                # Restore color and label logic
                if self.use_colors.get():
                    yticks = matrix_ax.get_yticklabels()
                    wrapped_labels = []
                    for label in yticks:
                        concept = label.get_text()
                        wrapped = wrap_label(concept, width=wrap_width)
                        wrapped_labels.append(wrapped)
                    matrix_ax.set_yticklabels(wrapped_labels)
                    for label, concept in zip(matrix_ax.get_yticklabels(), [l.get_text().replace('\n', ' ') for l in yticks]):
                        if concept in self.color_mapping:
                            color = self.color_mapping[concept]
                        else:
                            first_concept = concept.split('/')[0]
                            color = self.color_mapping.get(first_concept, 'black')
                        label.set_color(color)
                        label.set_weight('bold')
                        label.set_fontsize(10)
                else:
                    yticks = matrix_ax.get_yticklabels()
                    wrapped_labels = [wrap_label(label.get_text(), width=wrap_width) for label in yticks]
                    matrix_ax.set_yticklabels(wrapped_labels)
                    for label in matrix_ax.get_yticklabels():
                        label.set_color('black')
                        label.set_weight('normal')
                        label.set_fontsize(10)
                # Add parameter value labels in red
                for i, (bar, intersection) in enumerate(zip(bars, upset_index)):
                    if i == 0:
                        continue
                    if bar.get_height() == 0:
                        continue
                    x = bar.get_x() + bar.get_width() / 2
                    intersection_idx = i - 1
                    actual_intersection = upset_index[intersection_idx]
                    mask = np.ones(len(concept_matrix), dtype=bool)
                    for col, present in zip(concept_matrix_reset.columns, actual_intersection):
                        original_cols = col.split('/') if '/' in col else [col]
                        if present:
                            for orig_col in original_cols:
                                if orig_col in concept_matrix.columns:
                                    mask &= concept_matrix[orig_col] == 1
                        else:
                            for orig_col in original_cols:
                                if orig_col in concept_matrix.columns:
                                    mask &= concept_matrix[orig_col] == 0
                    param_vals = concept_matrix.loc[mask, varying_param].unique()
                    def fmt(v):
                        try:
                            f = float(v)
                            return f"{f:.1f}"
                        except Exception:
                            return str(v)
                    label = ','.join(fmt(v) for v in param_vals) if len(param_vals) > 0 else ''
                    print(f"Drawing parameter label: '{label}' at x={x}")
                    y_label = len(concept_matrix_reset.columns) - 0.3
                    matrix_ax.text(x, y_label, label, ha='center', va='bottom', fontsize=10, color='red', rotation=0, clip_on=False, weight='bold')
                # Add title and subtitle
                other_params = [p for p in params if p != varying_param]
                fixed_vals = {param_labels[p]: subset[p].iloc[0] for p in other_params}
                fixed_str = ' | '.join(f"{p}={v}" for p, v in fixed_vals.items())
                plt.suptitle(f"UpSet Diagram: {param_labels[varying_param]} sweep\nOther Params: {fixed_str}", fontsize=14, y=0.98)
                plt.subplots_adjust(top=0.88, bottom=0.12)
                plt.tight_layout(rect=[0, 0.12, 1, 0.88])
                outpath = f"{outdir}/compare_{param_labels[varying_param]}_composed.png"
                try:
                    plt.savefig(outpath, dpi=150, bbox_inches='tight', pad_inches=0.5, format='png')
                except Exception as e:
                    print(f"Error saving PNG diagram: {e}")
                plt.close('all')
                output_files.append(outpath)
            self.root.after(0, self.update_results, output_files)
            self.generate_html_table(blocks, df, self.color_mapping, outdir, param_labels)
            self.generate_docx_and_csv(blocks, df, self.color_mapping, outdir, param_labels)
            self.generate_stats_files(blocks, df, self.color_mapping, outdir, param_labels)
            self.root.after(0, self.display_block_info)
        except Exception as e:
            print('Exception in process_data:')
            traceback.print_exc()
            self.root.after(0, self.show_error, str(e))
    
    def update_results(self, output_files):
        # Get current window width
        window_width = self.root.winfo_width()
        
        # Update zoomable canvas with new images
        self.zoomable_canvas.set_images(output_files, window_width)
        
        # Stop progress and update status
        self.progress.stop()
        self.process_btn.config(state='normal')
        self.status_label.config(text=f"Processing complete. Generated {len(output_files)} plots.")
        
        # Call fit_to_window after processing
        self.zoomable_canvas.fit_to_window()
    
    def show_error(self, error_msg):
        self.progress.stop()
        self.process_btn.config(state='normal')
        self.status_label.config(text=f"Error occurred: {error_msg}")
        messagebox.showerror("Processing Error", f"An error occurred during processing:\n\n{error_msg}")

    def generate_html_table(self, blocks, df, color_mapping, outdir, param_labels):
        html_lines = [
            "<style>",
            ".concept-box { display: inline-block; padding: 1px 3px; margin: 1px; border-radius: 2px; font-weight: bold; white-space: nowrap; font-size: 10px; color: #fff; }",
            ".parameter-box { font-weight: bold; color: #333; font-size: 11px; }",
            ".varying-parameter { background-color: #ffffcc; }",
            ".concepts-line { white-space: nowrap; overflow-x: auto; }",
            "table { width: 100%; border-collapse: collapse; margin-bottom: 20px; font-size: 12px; }",
            "th, td { border: 1px solid #ddd; padding: 6px; text-align: left; vertical-align: top; }",
            "th { background-color: #f2f2f2; font-size: 11px; }",
            "</style>",
            "<table>",
            "<thead><tr>"
            "<th class='varying-parameter'>Temperature</th>"
            "<th class='varying-parameter'>Top P</th>"
            "<th class='varying-parameter'>Top K</th>"
            "<th class='varying-parameter'>BM25 Weight</th>"
            "<th>Extracted Concepts</th>"
            "</tr></thead><tbody>"
        ]
        for start, end, varying_param in blocks:
            subset = df.iloc[start:end].copy()
            for idx, row in subset.iterrows():
                html_lines.append("<tr>")
                html_lines.append(f"<td class='varying-parameter'><span class='parameter-box'>{row.get('Temperature','')}</span></td>")
                html_lines.append(f"<td class='varying-parameter'><span class='parameter-box'>{row.get('Top-p','')}</span></td>")
                html_lines.append(f"<td class='varying-parameter'><span class='parameter-box'>{row.get('Top-k','')}</span></td>")
                html_lines.append(f"<td class='varying-parameter'><span class='parameter-box'>{row.get('BM25 Weight','')}</span></td>")
                html_lines.append("<td><div class='concepts-line'>")
                for concept in row['Concepts']:
                    color = color_mapping.get(concept, '#888')
                    html_lines.append(
                        f"<span class='concept-box' style='background-color: {color};'>{html.escape(concept)}</span>"
                    )
                html_lines.append("</div></td></tr>")
        html_lines.append("</tbody></table>")
        with open(os.path.join(outdir, "compare.htm"), "w", encoding="utf-8") as f:
            f.write("\n".join(html_lines))

    def get_wrap_width(self):
        try:
            return int(self.wrap_width_var.get())
        except Exception:
            return 75

    def merge_csv_files(self):
        import re
        import traceback
        file_paths = filedialog.askopenfilenames(
            title="Select CSV files to merge",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        if not file_paths:
            return
        # Store the folder name for DOCX title
        folder_name = os.path.basename(os.path.dirname(file_paths[0])) if file_paths else ""
        self.merged_folder_name = folder_name
        self.merge_label.config(text="Selected files:\n" + "\n".join([os.path.basename(fp) for fp in file_paths]), foreground="black")
        error_patterns = [
            r"error generating response:",
            r"api error occurred:",
            r"bad gateway",
            r"cloudflare",
            r"server disconnected without sending a response",
            r"getaddrinfo failed"
        ]
        error_found = False
        error_msgs = []
        dfs = []
        non_utf8_files = []
        for fp in file_paths:
            df = None
            last_exc = None
            used_encoding = None
            try:
                df = pd.read_csv(fp, encoding='utf-8')
                used_encoding = 'utf-8'
            except Exception as e:
                last_exc = e
                non_utf8_files.append(os.path.basename(fp))
                print(f"File {os.path.basename(fp)} could not be read as utf-8: {e}")
            if used_encoding and used_encoding != 'utf-8':
                print(f"Warning: File {os.path.basename(fp)} was read with encoding {used_encoding}, not utf-8. This may cause issues with special characters.")
            if df is None:
                error_found = True
                error_msgs.append(f"File: {os.path.basename(fp)}, Error reading file as utf-8: {last_exc}")
                continue
            for idx, row in df.iterrows():
                for col in df.columns:
                    val = str(row[col]).lower()
                    for pat in error_patterns:
                        if re.search(pat, val):
                            error_found = True
                            error_msgs.append(f"File: {os.path.basename(fp)}, Row: {idx+2}, Column: '{col}', Error: {row[col]}")
            dfs.append(df)
        if non_utf8_files:
            msg = "Merge aborted: The following files are not valid UTF-8 and may cause encoding issues. Please convert them to UTF-8 and try again:\n" + "\n".join(non_utf8_files)
            self.merge_label.config(text=msg, foreground="red")
            messagebox.showerror("Merge Error", msg)
            return
        if error_found:
            msg = "Merge aborted due to invalid data in the following locations:\n" + "\n".join(error_msgs)
            self.merge_label.config(text=msg, foreground="red")
            messagebox.showerror("Merge Error", msg)
            return
        try:
            merged_df = pd.concat(dfs, ignore_index=True)
            merged_path = os.path.join(os.getcwd(), "compare_input.csv")
            merged_df.to_csv(merged_path, index=False, encoding='utf-8')
            self.csv_file = merged_path
            self.file_label.config(text=os.path.basename(merged_path), foreground="black")
            self.merge_label.config(text=f"Merged:\n" + "\n".join([os.path.basename(fp) for fp in file_paths]) + f"\n(encoding: utf-8)", foreground="black")
            self.process_btn.config(state='normal')
        except Exception as e:
            import sys
            import traceback
            tb_str = traceback.format_exc()
            print(f"Failed to merge files: {e}\nTraceback:\n{tb_str}")
            error_detail = f"Failed to merge files: {e}\nFiles attempted: {', '.join([os.path.basename(fp) for fp in file_paths])}\nTraceback (see shell):\n{tb_str}"
            self.merge_label.config(text=error_detail, foreground="red")
            messagebox.showerror("Merge Error", error_detail)

    def generate_docx_and_csv(self, blocks, df, color_mapping, outdir, param_labels):
        try:
            import docx
            from docx.shared import RGBColor
            import csv
            doc = docx.Document()
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Temperature'
            hdr_cells[1].text = 'Top P'
            hdr_cells[2].text = 'Top K'
            hdr_cells[3].text = 'BM25 Weight'
            hdr_cells[4].text = 'Extracted Concepts'
            for start, end, varying_param in blocks:
                subset = df.iloc[start:end].copy()
                for idx, row in subset.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row.get('Temperature',''))
                    row_cells[1].text = str(row.get('Top-p',''))
                    row_cells[2].text = str(row.get('Top-k',''))
                    row_cells[3].text = str(row.get('BM25 Weight',''))
                    concept_strs = []
                    for concept in row['Concepts']:
                        color = color_mapping.get(concept, '#888')
                        concept_strs.append(f"{concept} [{color}]")
                    row_cells[4].text = ", ".join(concept_strs)
            docx_path = os.path.join(outdir, "compare.docx")
            doc.save(docx_path)
            csv_path = os.path.join(outdir, "compare.csv")
            with open(csv_path, "w", newline='', encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(['Temperature', 'Top P', 'Top K', 'BM25 Weight', 'Extracted Concepts'])
                for start, end, varying_param in blocks:
                    subset = df.iloc[start:end].copy()
                    for idx, row in subset.iterrows():
                        concept_strs = []
                        for concept in row['Concepts']:
                            color = color_mapping.get(concept, '#888')
                            concept_strs.append(f"{concept} [{color}]")
                        writer.writerow([
                            row.get('Temperature',''),
                            row.get('Top-p',''),
                            row.get('Top-k',''),
                            row.get('BM25 Weight',''),
                            ", ".join(concept_strs)
                        ])
        except Exception as e:
            print(f"Error generating DOCX/CSV: {e}")

    def generate_stats_files(self, blocks, df, color_mapping, outdir, param_labels):
        try:
            import pandas as pd
            import docx
            from docx.shared import RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_ORIENT
            from docx2pdf import convert as docx2pdf_convert
            import os
            txt_lines = []
            all_concepts = sorted(set(c for concepts in df['Concepts'] for c in concepts))
            # TXT
            txt_lines.append("Concepts Overview\n================\n")
            txt_lines.append(f"Total unique concepts: {len(all_concepts)}\n")
            txt_lines.append("\nAll Concepts:\n" + "\n".join(all_concepts) + "\n")
            txt_lines.append("\nConcepts per block:\n")
            for start, end, varying_param in blocks:
                subset = df.iloc[start:end].copy()
                block_concepts = sorted(set(c for concepts in subset['Concepts'] for c in concepts))
                txt_lines.append(f"{varying_param} ({start}-{end-1}): {len(block_concepts)} concepts\n" + ", ".join(block_concepts) + "\n")
            txt_lines.append("\nColor Groups:\n")
            color_groups = {}
            for concept, color in color_mapping.items():
                color_groups.setdefault(color, []).append(concept)
            for color, concepts in color_groups.items():
                txt_lines.append(f"Color {color}: {len(concepts)} concepts\n" + ", ".join(concepts) + "\n")
            txt_path = os.path.join(outdir, "compare_stats.txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write("".join(txt_lines))
            # DOCX
            doc = docx.Document()
            section = doc.sections[0]
            from docx.shared import Inches
            from docx.enum.section import WD_ORIENT
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap width and height for landscape
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            # Set narrow margins and header/footer
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.3)
            section.left_margin = Inches(0.3)
            section.right_margin = Inches(0.3)
            section.header_distance = Inches(0.3)
            section.footer_distance = Inches(0.3)
            # --- First page: folder name as title ---
            folder_name = getattr(self, 'merged_folder_name', '')
            if folder_name:
                doc.add_heading(folder_name, level=1)
            else:
                doc.add_heading("Merged Results", level=1)
            # Insert 4 images in a table, side by side, width matches table width
            image_dir = os.path.join(os.path.dirname(outdir), "compare_gui_output")
            image_files = [
                "compare_BM25_composed.png",
                "compare_Topk_composed.png",
                "compare_Topp_composed.png",
                "compare_Temp_composed.png"
            ]
            img_table = doc.add_table(rows=1, cols=4)
            img_table.autofit = False
            # Calculate available width and height for images (page width - margins)
            table_width = int(section.page_width - section.left_margin - section.right_margin)
            available_height = int(section.page_height - section.top_margin - section.bottom_margin)
            img_width = int(table_width // 4)
            from PIL import Image as PILImage
            for i, img_file in enumerate(image_files):
                img_path = os.path.join(image_dir, img_file)
                if os.path.exists(img_path):
                    cell = img_table.rows[0].cells[i]
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cell.paragraphs[0].add_run()
                    # Open image to get aspect ratio
                    with PILImage.open(img_path) as pil_img:
                        aspect = pil_img.width / pil_img.height
                        # Calculate height in inches to fill page height
                        height_in_inches = (available_height / 914400)
                        height_in_inches = max(0.5, height_in_inches - 0.1)
                        # Calculate width to maintain aspect ratio, but cap at img_width
                        width_in_inches = min(img_width / 914400, aspect * height_in_inches)
                        run.add_picture(img_path, width=docx.shared.Inches(width_in_inches))
            # --- Page break, then Color Groups section ---
            doc.add_heading("Color Groups", level=1)
            # Color Groups table: 4 columns, set widths proportionally to fill table_width
            color_table = doc.add_table(rows=1, cols=4)
            color_table.autofit = False
            col_props = [1.2, 0.4, 2.2, 3.2]
            total = sum(col_props)
            col_widths = [int(table_width * (w / total)) for w in col_props]
            for i, w in enumerate(col_widths):
                color_table.columns[i].width = w
            color_table.rows[0].cells[0].text = 'Group Label'
            color_table.rows[0].cells[1].text = 'Count'
            color_table.rows[0].cells[2].text = 'Concepts'
            color_table.rows[0].cells[3].text = 'Parameters'
            for color, concepts in color_groups.items():
                # Group label: common word in group (use the longest common substring or first word)
                group_label = ''
                if len(concepts) > 1:
                    from difflib import SequenceMatcher
                    def lcs(a, b):
                        match = SequenceMatcher(None, a, b).find_longest_match(0, len(a), 0, len(b))
                        return a[match.a: match.a + match.size]
                    lcs_str = concepts[0]
                    for c in concepts[1:]:
                        lcs_str = lcs(lcs_str, c)
                    group_label = lcs_str.strip() if lcs_str.strip() else concepts[0].split()[0]
                else:
                    group_label = concepts[0].split()[0]
                row_cells = color_table.add_row().cells
                row_cells[0].text = group_label
                row_cells[1].text = str(len(concepts))
                # Concepts (colored)
                para = row_cells[2].paragraphs[0]
                for i, concept in enumerate(concepts):
                    run = para.add_run(concept)
                    if color.startswith('#') and len(color) == 7:
                        r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                        run.font.color.rgb = RGBColor(r, g, b)
                    if i < len(concepts) - 1:
                        para.add_run(", ")
                # Parameters (expanded cell)
                param_strs = []
                for concept in concepts:
                    param_rows = df[df['Concepts'].apply(lambda lst: concept in lst)]
                    for _, row in param_rows.iterrows():
                        param_strs.append(f"(Temp: {row.get('Temperature','')}, Topp: {row.get('Top-p','')}, Topk: {row.get('Top-k','')}, BM25: {row.get('BM25 Weight','')})")
                row_cells[3].text = ", ".join(param_strs)
                # Color group label cell colored
                for para in row_cells[0].paragraphs:
                    for run in para.runs:
                        if color.startswith('#') and len(color) == 7:
                            r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                            run.font.color.rgb = RGBColor(r, g, b)
            # --- All Concepts Table ---
            doc.add_heading("All Concepts", level=1)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Index'
            hdr_cells[1].text = 'Concept'
            for idx, concept in enumerate(all_concepts):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx+1)
                para = row_cells[1].paragraphs[0]
                run = para.add_run(concept)
                color = color_mapping.get(concept, '#888')
                if color.startswith('#') and len(color) == 7:
                    r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                    run.font.color.rgb = RGBColor(r, g, b)
            # --- Concepts per Block Table ---
            doc.add_heading("Concepts per Block", level=1)
            block_table = doc.add_table(rows=1, cols=3)
            block_table.autofit = True
            block_table.rows[0].cells[0].text = 'Block'
            block_table.rows[0].cells[1].text = 'Concepts'
            block_table.rows[0].cells[2].text = 'Count'
            for start, end, varying_param in blocks:
                subset = df.iloc[start:end].copy()
                block_concepts = sorted(set(c for concepts in subset['Concepts'] for c in concepts))
                row_cells = block_table.add_row().cells
                row_cells[0].text = varying_param
                para = row_cells[1].paragraphs[0]
                for i, concept in enumerate(block_concepts):
                    run = para.add_run(concept)
                    color = color_mapping.get(concept, '#888')
                    if color.startswith('#') and len(color) == 7:
                        r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                        run.font.color.rgb = RGBColor(r, g, b)
                    if i < len(block_concepts) - 1:
                        para.add_run(", ")
                row_cells[2].text = str(len(block_concepts))
            # --- UpSet Histogram Section ---
            doc.add_heading("UpSet Diagram Histograms", level=1)
            # For each block, add a table with parameter, value, and histogram (bar heights/counts)
            for start, end, varying_param in blocks:
                subset = df.iloc[start:end].copy()
                doc.add_heading(f"{varying_param} Sweep", level=2)
                hist_table = doc.add_table(rows=1, cols=3)
                hist_table.autofit = True
                hist_table.rows[0].cells[0].text = 'Parameter'
                hist_table.rows[0].cells[1].text = 'Value'
                hist_table.rows[0].cells[2].text = 'Histogram (bar counts)'
                # For each unique value of the varying parameter, count the number of concepts present
                for val in sorted(subset[varying_param].unique()):
                    mask = subset[varying_param] == val
                    present_concepts = [c for concepts in subset[mask]['Concepts'] for c in concepts]
                    hist_row = hist_table.add_row().cells
                    hist_row[0].text = varying_param
                    hist_row[1].text = str(val)
                    hist_row[2].text = str(len(present_concepts))
            docx_path = os.path.join(outdir, "compare_stats.docx")
            doc.save(docx_path)
            # PDF: convert DOCX to PDF
            LIBREOFFICE_PATH = r"C:\\Program Files\\LibreOffice\\program\\soffice.exe"
            result = subprocess.run([
                LIBREOFFICE_PATH,
                "--headless",
                "--convert-to", "pdf",
                docx_path,
                "--outdir", outdir
            ], check=True, capture_output=True, text=True)
            pdf_path = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
            if not os.path.exists(pdf_path):
                raise FileNotFoundError(f"PDF not found at {pdf_path}")
            self.merge_label.config(text="Stats PDF generated successfully.", foreground="black")
        except Exception as e:
            print(f"Error generating stats TXT/PDF/DOCX: {e}")

    def on_group_by_subletters(self):
        if self.group_by_subletters.get():
            self.group_by_words.set(False)

    def on_group_by_words(self):
        if self.group_by_words.get():
            self.group_by_subletters.set(False)

    def on_color_criteria_toggle(self):
        if self.use_colors.get():
            self.group_by_same_color_cb.state(['!disabled'])
        else:
            self.group_by_same_color_cb.state(['disabled'])
            self.group_by_same_color.set(False)

    def on_agg_group_by_subletters(self):
        if self.agg_group_by_subletters.get():
            self.agg_group_by_words.set(False)

    def on_agg_group_by_words(self):
        if self.agg_group_by_words.get():
            self.agg_group_by_subletters.set(False)

    def on_agg_color_criteria_toggle(self):
        # This method can be used for future functionality if needed
        pass

    def on_agg_enable_fuzzy_toggle(self):
        # Enable/disable fuzzy logic controls based on checkbox
        if self.agg_enable_fuzzy.get():
            self.sim_threshold_entry.config(state='normal')
            self.grouping_logic_combo.config(state='readonly')
        else:
            self.sim_threshold_entry.config(state='disabled')
            self.grouping_logic_combo.config(state='disabled')

    def get_aggregation_params(self):
        """Get aggregation parameters from GUI controls"""
        try:
            threshold = float(self.sim_threshold_var.get())
        except Exception:
            threshold = 0.85
        
        grouping_logic = self.grouping_logic_var.get()
        if not self.agg_enable_fuzzy.get():
            grouping_logic = "None"
        
        try:
            min_letters = int(self.agg_min_letters_var.get())
        except Exception:
            min_letters = 5
        
        return {
            'threshold': threshold,
            'grouping_logic': grouping_logic,
            'min_letters': min_letters,
            'use_colors': self.agg_use_colors.get(),
            'group_by_words': self.agg_group_by_words.get(),
            'group_by_subletters': self.agg_group_by_subletters.get()
        }

    def select_aggregate_folder(self):
        folder_path = filedialog.askdirectory(title="Select Parent Folder Containing Results")
        if folder_path:
            self.aggregate_folder = folder_path
            self.aggregate_folder_label.config(text=folder_path, foreground="black")
            self.aggregate_btn.config(state='normal')
            self.aggregate_status_label.config(text="Ready to aggregate.", foreground="blue")
        else:
            self.aggregate_folder = None
            self.aggregate_folder_label.config(text="No folder selected", foreground="gray")
            self.aggregate_btn.config(state='disabled')
            self.aggregate_status_label.config(text="", foreground="blue")

    def aggregate_results(self):
        # Run aggregation in a background thread to avoid blocking the UI
        thread = threading.Thread(target=self._aggregate_results_thread)
        thread.daemon = True
        thread.start()

    def _aggregate_results_thread(self):
        import os
        import docx
        from docx.shared import RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from collections import defaultdict, Counter
        import re
        from tkinter import messagebox
        from PIL import Image as PILImage
        import difflib
        import traceback
        import time
        import string
        def safe_update_progress(val, elapsed, est_total):
            self.root.after(0, lambda: self._update_aggregate_progress(val, elapsed, est_total))
        def safe_update_status(msg, color):
            self.root.after(0, lambda: self.aggregate_status_label.config(text=msg, foreground=color))
        def safe_update_time(msg):
            self.root.after(0, lambda: self.aggregate_time_label.config(text=msg))
        safe_update_status("Scanning folders...", "blue")
        parent = self.aggregate_folder
        if not parent or not os.path.isdir(parent):
            safe_update_status("Invalid parent folder.", "red")
            return
        # Get aggregation parameters from GUI
        params = self.get_aggregation_params()
        threshold = params['threshold']
        grouping_logic = params['grouping_logic']
        # Step 1: Find valid subfolders
        subfolders = [os.path.join(parent, d) for d in os.listdir(parent) if os.path.isdir(os.path.join(parent, d))]
        valid_folders = []
        for folder in subfolders:
            files = os.listdir(folder)
            required = ["compare_stats.txt", "compare_BM25_composed.png", "compare_Topk_composed.png", "compare_Topp_composed.png", "compare_Temp_composed.png"]
            if all(f in files for f in required):
                valid_folders.append(folder)
        if not valid_folders:
            safe_update_status("No valid result folders found.", "red")
            return
        safe_update_status(f"Found {len(valid_folders)} valid folders. Extracting data...", "blue")
        # Step 2: Extract concepts and images
        folder_concepts = {}
        all_concepts = set()
        t0 = time.time()
        for idx, folder in enumerate(valid_folders):
            txt_path = os.path.join(folder, "compare_stats.txt")
            try:
                with open(txt_path, encoding="utf-8") as f:
                    lines = f.readlines()
                # Find 'All Concepts:' section
                concepts = []
                in_concepts = False
                for line in lines:
                    if line.strip().startswith("All Concepts:"):
                        in_concepts = True
                        continue
                    if in_concepts:
                        if not line.strip() or line.strip().endswith(":"):
                            break
                        c = line.strip()
                        if c:
                            concepts.append(c)
                folder_concepts[folder] = set(concepts)
                all_concepts.update(concepts)
            except Exception as e:
                print(f"Error reading {txt_path}: {e}")
            # Progress bar update
            elapsed = time.time() - t0
            if idx == 0 and len(valid_folders) > 1:
                est_total = elapsed * len(valid_folders)
            else:
                est_total = elapsed if idx == 0 else est_total
            safe_update_progress(100 * (idx+1) / len(valid_folders), elapsed, est_total)
        total_elapsed = time.time() - t0
        # Step 3: Refined Color Grouping Logic
        concept_list = sorted(all_concepts, key=lambda x: x.lower())
        # LLM Grouping logic
        if self.llm_grouping_var.get():
            # Improved prompt for group-name: [concepts] format
            prompt = (
                self.llm_prompt_var.get().strip() +
                "\nReturn a JSON object where each key is a group name and the value is a list of concepts. Example:\n{\n  \"Character\": [\"Character (1103a17–1103b35)\", ...],\n  \"Intellect\": [\"Intellect (Nous)\", ...]\n}"
            )
            model = self.llm_model_var.get()
            llm_groups = real_llm_grouping(concept_list, prompt, model, output_dir=parent)
            # Only assign colors to concepts present in LLM output
            concepts_in_llm = set()
            for group in llm_groups:
                for concept in group:
                    concepts_in_llm.add(concept)
            missing_concepts = [c for c in concept_list if c not in concepts_in_llm]
            if missing_concepts:
                msg = f"[LLM GROUPING] {len(missing_concepts)} concepts missing from LLM output. Not assigned to any group.\n" + ", ".join(missing_concepts)
                print(msg)
                self.root.after(0, lambda: self.aggregate_status_label.config(text=msg, foreground="red"))
            color_palette = [
                '#800000', '#FF8C00', '#228B22', '#8B008B', '#A0522D', '#2E8B57', '#9932CC', '#FFD700',
                '#556B2F', '#C71585', '#8B4513', '#20B2AA', '#B22222', '#FF4500', '#6A5ACD', '#D2691E',
                '#006400', '#708090', '#FF6347', '#483D8B', '#000000', '#808000', '#8B0000', '#FF1493',
            ]
            group_colors = {}
            color_to_concepts = defaultdict(list)
            for i, group in enumerate(llm_groups):
                color = color_palette[i % len(color_palette)]
                for concept in group:
                    group_colors[concept] = color
                    color_to_concepts[color].append(concept)

            # For LLM grouping, preserve group structure and color for DOCX output
            if self.llm_grouping_var.get():
                llm_group_tuples = []
                for i, group in enumerate(llm_groups):
                    color = color_palette[i % len(color_palette)]
                    llm_group_tuples.append((color, group))
        else:
            def create_refined_color_groups(concepts, use_colors, group_by_words, group_by_subletters, min_letters):
                """Create refined color groups based on semantic similarity (no transitive merging)"""
                if not use_colors:
                    return {c: '#000000' for c in concepts}
                stopwords = set(['the','a','an','and','or','but','in','on','at','to','for','of','with','by','is','are','was','were','be','been','being','have','has','had','do','does','did','will','would','should','could','can','may','might','must'])
                def extract_words(concept):
                    words = re.findall(r'\b\w+\b', concept.lower())
                    return [w for w in words if w not in stopwords and len(w) > 2]
                # Build word-to-concept mapping
                word_to_concepts = {}
                for concept in concepts:
                    words = extract_words(concept)
                    for w in words:
                        word_to_concepts.setdefault(w, set()).add(concept)
                # Build groups: each group is all concepts sharing a word
                groups = []
                assigned = set()
                for w, cset in word_to_concepts.items():
                    group = set(cset) - assigned
                    if len(group) > 1:
                        groups.append(group)
                        assigned.update(group)
                # Any unassigned concepts become their own group
                for concept in concepts:
                    if concept not in assigned:
                        groups.append({concept})
                # Assign colors
                color_palette = [
                    '#800000', '#FF8C00', '#228B22', '#8B008B', '#A0522D', '#2E8B57', '#9932CC', '#FFD700',
                    '#556B2F', '#C71585', '#8B4513', '#20B2AA', '#B22222', '#FF4500', '#6A5ACD', '#D2691E',
                    '#006400', '#708090', '#FF6347', '#483D8B', '#000000', '#808000', '#8B0000', '#FF1493',
                ]
                group_colors = {}
                for i, group in enumerate(groups):
                    color = color_palette[i % len(color_palette)]
                    for concept in group:
                        group_colors[concept] = color
                return group_colors
            group_colors = create_refined_color_groups(
                concept_list, 
                params['use_colors'],
                params['group_by_words'],
                params['group_by_subletters'],
                params['min_letters']
            )
            # --- Fuzzy/Exact/None grouping for overlap and canonicalization ---
            if self.agg_enable_fuzzy.get() and grouping_logic == "Fuzzy":
                groups = []  # List of sets
                used = set()
                for c in concept_list:
                    if c in used:
                        continue
                    group = set([c])
                    for other in concept_list:
                        if other == c or other in used:
                            continue
                        ratio = difflib.SequenceMatcher(None, c.lower(), other.lower()).ratio()
                        if ratio >= threshold:
                            group.add(other)
                            used.add(other)
                    used.add(c)
                    groups.append(group)
            elif grouping_logic == "Exact":
                groups = [{c} for c in concept_list]
            else:  # None or fuzzy disabled
                groups = [{c} for c in concept_list]
            
            # Map each concept to its canonical group label (first in sorted group)
            group_labels = {}
            group_variants = {}
            for group in groups:
                label = sorted(group, key=lambda x: x.lower())[0]
                for variant in group:
                    group_labels[variant] = label
                group_variants[label] = sorted(group, key=lambda x: x.lower())
            canonical_concepts = sorted(group_variants.keys(), key=lambda x: x.lower())
            
            # Step 4: Compute overlap/uniqueness using color groups (unique semantic groups)
            color_to_concepts = defaultdict(list)
            for c in concept_list:
                color = group_colors.get(c, '#FF0000')
                if c not in group_colors:
                    print(f"[GROUP COLOR WARNING] Concept not in group_colors: {c}")
                    self.root.after(0, lambda c=c: self.aggregate_status_label.config(text=f"[GROUP COLOR WARNING] Concept not in group_colors: {c}", foreground='red'))
                    group_colors[c] = '#FF0000'
                color_to_concepts[color].append(c)
        # else: color_to_concepts is already correct from LLM grouping
        
        # Create unique color groups (each color represents a unique semantic group)
        unique_color_groups = list(color_to_concepts.keys())
        
        # Compute overlap using color groups instead of canonical concepts
        color_overlap_table = []  # List of (color, [present in folder1, folder2, ...])
        folder_names = [os.path.basename(f) for f in valid_folders]
        for color in unique_color_groups:
            row = [color]
            concepts_in_color = set(color_to_concepts[color])
            for folder in valid_folders:
                present = any(c in folder_concepts[folder] for c in concepts_in_color)
                row.append(1 if present else 0)
            color_overlap_table.append(row)
        
        # Keep the original canonical overlap table for the final table
        if not self.llm_grouping_var.get():
            overlap_table = []  # List of (canonical, [present in folder1, folder2, ...])
            for canon in canonical_concepts:
                row = [canon]
                canon_variants = set(group_variants[canon])
                for folder in valid_folders:
                    present = any(v in folder_concepts[folder] for v in canon_variants)
                    row.append(1 if present else 0)
                overlap_table.append(row)
        # --- Unique words in folder names ---
        def get_unique_folder_words(folder_names):
            all_words = [set(re.findall(r'\w+', name.lower())) for name in folder_names]
            common = set.intersection(*all_words) if all_words else set()
            unique_words = []
            for words in all_words:
                unique = words - common
                unique_words.append(unique)
            return unique_words, common
        unique_folder_words, common_folder_words = get_unique_folder_words(folder_names)
        folder_unique_map = {folder_names[i]: unique_folder_words[i] for i in range(len(folder_names))}
        def get_concept_folders(concept, folder_concepts, folder_names, valid_folders):
            folders = set()
            for i, folder in enumerate(valid_folders):
                if concept in folder_concepts[folder]:
                    folders.add(folder_names[i])
            return folders
        def get_concept_unique_words(concept, folder_concepts, folder_names, valid_folders, folder_unique_map):
            folders = get_concept_folders(concept, folder_concepts, folder_names, valid_folders)
            unique_words = set()
            for fname in folders:
                unique_words |= folder_unique_map[fname]
            return unique_words
        # Step 5: Collect images
        folder_images = []  # List of (folder_name, [img1, img2, img3, img4])
        for folder in valid_folders:
            imgs = [os.path.join(folder, f) for f in ["compare_BM25_composed.png", "compare_Topk_composed.png", "compare_Topp_composed.png", "compare_Temp_composed.png"]]
            folder_images.append((os.path.basename(folder), imgs))
        # Step 6: Generate DOCX
        try:
            doc = docx.Document()
            # Set landscape orientation and zero margins
            section = doc.sections[0]
            from docx.enum.section import WD_ORIENT
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
            section.top_margin = 0
            section.bottom_margin = 0
            section.left_margin = 0
            section.right_margin = 0
            section.header_distance = 0
            section.footer_distance = 0

            # Gather folder/group stats for the left column text
            total_concepts = sum(len(group) for _, group in llm_group_tuples) if self.llm_grouping_var.get() else sum(len(concepts) for concepts in color_to_concepts.values())
            folder_concept_counts = []
            folder_group_counts = []
            for folder in valid_folders:
                concepts_in_folder = set()
                groups_in_folder = 0
                if self.llm_grouping_var.get():
                    for _, group in llm_group_tuples:
                        if any(concept in folder_concepts[folder] for concept in group):
                            groups_in_folder += 1
                            concepts_in_folder.update([concept for concept in group if concept in folder_concepts[folder]])
                else:
                    for color, concepts in color_to_concepts.items():
                        if any(concept in folder_concepts[folder] for concept in concepts):
                            groups_in_folder += 1
                            concepts_in_folder.update([concept for concept in concepts if concept in folder_concepts[folder]])
                folder_concept_counts.append(len(concepts_in_folder))
                folder_group_counts.append(groups_in_folder)

            # Regrouping method
            if self.llm_grouping_var.get():
                method_str = "color grouping based on LLM"
            else:
                if params['group_by_words']:
                    method_str = "color grouping by words"
                elif params['group_by_subletters']:
                    method_str = "color grouping by subletters"
                elif self.agg_enable_fuzzy.get() and params['grouping_logic'] == "Fuzzy":
                    method_str = f"fuzzy grouping (threshold {params['threshold']})"
                else:
                    method_str = "exact grouping"
            num_groups = len(llm_group_tuples) if self.llm_grouping_var.get() else len(color_to_concepts)
            total_unique_groups = len(unique_color_groups)
            
            # Function to extract group name from concepts
            def extract_group_name(concepts):
                """Extract a single common word that represents the group"""
                if not concepts:
                    return "Unknown"
                
                # Try to find the most common meaningful word
                all_words = []
                for concept in concepts:
                    words = re.findall(r'\b\w+\b', concept.lower())
                    # Filter out common words and short words
                    meaningful_words = [w for w in words if len(w) > 3 and w not in ['the', 'and', 'for', 'with', 'from', 'that', 'this', 'have', 'been', 'they', 'will', 'would', 'could', 'should']]
                    all_words.extend(meaningful_words)
                
                if all_words:
                    # Count word frequencies
                    from collections import Counter
                    word_counts = Counter(all_words)
                    # Return the most common word
                    return word_counts.most_common(1)[0][0].capitalize()
                
                # Fallback: use first word of first concept
                first_concept = concepts[0]
                words = re.findall(r'\b\w+\b', first_concept)
                if words:
                    return words[0].capitalize()
                
                return "Group"
            
            # Create data for UpSet diagram
            group_data = []
            group_names = []
            
            if self.llm_grouping_var.get():
                for color, concepts in llm_group_tuples:
                    group_name = extract_group_name(concepts)
                    group_names.append(group_name)
                    
                    # Create row for this group: [group_name, present_in_folder1, present_in_folder2, ...]
                    row = [group_name]
                    concepts_in_color = set(concepts)
                    for folder in valid_folders:
                        present = any(c in folder_concepts[folder] for c in concepts_in_color)
                        row.append(1 if present else 0)
                    group_data.append(row)
            else:
                for color in unique_color_groups:
                    concepts = color_to_concepts[color]
                    group_name = extract_group_name(concepts)
                    group_names.append(group_name)
                    
                    # Create row for this group
                    row = [group_name]
                    concepts_in_color = set(concepts)
                    for folder in valid_folders:
                        present = any(c in folder_concepts[folder] for c in concepts_in_color)
                        row.append(1 if present else 0)
                    group_data.append(row)
            
            # Create DataFrame for UpSet
            import pandas as pd
            from upsetplot import UpSet, from_indicators
            import matplotlib.pyplot as plt
            
            # Create DataFrame with group names as COLUMNS (not index) - like the existing code
            # This matches the pattern: concept_matrix_reset with concepts as columns
            df_groups = pd.DataFrame(group_data, columns=['Group'] + folder_names)
            
            # Ensure unique group names by adding index if duplicates exist
            unique_group_names = []
            group_name_counts = {}
            for group_name in df_groups['Group']:
                if group_name in group_name_counts:
                    group_name_counts[group_name] += 1
                    unique_name = f"{group_name}_{group_name_counts[group_name]}"
                else:
                    group_name_counts[group_name] = 0
                    unique_name = group_name
                unique_group_names.append(unique_name)
            
            df_groups['Group'] = unique_group_names
            
            # Now create the final DataFrame with groups as columns (like concept_matrix_reset)
            # Extract the boolean data (folder columns) and use group names as column headers
            group_columns = df_groups[folder_names].T  # Transpose to get groups as columns
            group_columns.columns = unique_group_names  # Set group names as column headers
            
            # Convert to boolean and reset index to get default integer index
            df_for_upset = group_columns.astype(bool).reset_index(drop=True)
            
            print(f"[UPSET DEBUG] Final DataFrame shape: {df_for_upset.shape}")
            print(f"[UPSET DEBUG] Final DataFrame columns: {df_for_upset.columns.tolist()}")
            print(f"[UPSET DEBUG] Final DataFrame dtypes: {df_for_upset.dtypes}")
            print(f"[UPSET DEBUG] Final DataFrame head:")
            print(df_for_upset.head())
            
            # Create UpSet plot with error handling
            try:
                # Use the same pattern as the existing code
                upset_data = from_indicators(df_for_upset, df_for_upset.columns)
                print('[UPSET DEBUG] UpSet data created successfully')
                
                # Create figure that fits page height (landscape orientation)
                # Calculate available height for the plot
                page_height_inches = 8.5  # Landscape page height
                margin_inches = 1.0  # Top and bottom margins
                available_height = page_height_inches - (2 * margin_inches)
                
                fig, axes = plt.subplots(1, 1, figsize=(12, available_height))
                
                upset = UpSet(upset_data, show_counts=True)
                axes = upset.plot(fig=fig)
                bar_ax = axes['intersections']
                matrix_ax = axes['matrix']
                bars = bar_ax.patches
                upset_index = upset_data.index
                
                # Add group colors to the y-axis labels (group names)
                yticks = matrix_ax.get_yticklabels()
                for label, group_name in zip(matrix_ax.get_yticklabels(), df_for_upset.columns):
                    # Find the color for this group
                    group_color = None
                    if self.llm_grouping_var.get():
                        # For LLM grouping, find the color from llm_group_tuples
                        for color, concepts in llm_group_tuples:
                            if extract_group_name(concepts) == group_name:
                                group_color = color
                                break
                    else:
                        # For color grouping, find the color from color_to_concepts
                        for color, concepts in color_to_concepts.items():
                            if extract_group_name(concepts) == group_name:
                                group_color = color
                                break
                    
                    if group_color and group_color.startswith('#') and len(group_color) == 7:
                        label.set_color(group_color)
                    else:
                        label.set_color('black')
                    label.set_weight('bold')
                    label.set_fontsize(10)
                
                # Add short folder names as red column labels (rotated 90 degrees)
                # Fixed logic: Skip first bar but ensure labels match correct intersections
                
                # First, collect all non-empty bars with their indices
                non_empty_bars = []
                for i, bar in enumerate(bars):
                    if bar.get_height() > 0:
                        non_empty_bars.append((i, bar))
                
                print(f"[UPSET DEBUG] Total bars: {len(bars)}, Non-empty bars: {len(non_empty_bars)}")
                
                # Create a list of all folder names to cycle through
                all_folder_names = []
                for folder in folder_names:
                    folder_unique_words = folder_unique_map.get(folder, set())
                    if folder_unique_words:
                        short_name = sorted(folder_unique_words)[0].capitalize()
                    else:
                        short_name = folder.split()[0].capitalize()
                    all_folder_names.append(short_name)
                
                # Reverse to match the original logic
                all_folder_names = all_folder_names[::-1]
                
                print(f"[UPSET DEBUG] All folder names for cycling: {all_folder_names}")
                
                # Process all non-empty bars except the first one (skip index 0)
                print(f"[UPSET DEBUG] Non-empty bars indices: {[b[0] for b in non_empty_bars]}")
                
                # Filter out the first bar (index 0) and process the rest
                bars_to_process = [(idx, bar) for idx, bar in non_empty_bars if idx != 0]
                print(f"[UPSET DEBUG] Bars to process: {[b[0] for b in bars_to_process]}")
                
                for bar_idx, (original_index, bar) in enumerate(bars_to_process):
                    print(f"[UPSET DEBUG] Processing bar_idx={bar_idx}, original_index={original_index}")
                    x = bar.get_x() + bar.get_width() / 2
                    
                    # Get the intersection data for this specific bar using its original index
                    actual_intersection = upset_index[original_index]
                    
                    # Find which folders are present in this intersection
                    present_folders = []
                    for j, present in enumerate(actual_intersection):
                        if present and j < len(folder_names):
                            # Get the short folder name from the unique words
                            folder_unique_words = folder_unique_map.get(folder_names[j], set())
                            if folder_unique_words:
                                # Use the first unique word as short name
                                short_name = sorted(folder_unique_words)[0].capitalize()
                            else:
                                # Fallback to first word of folder name
                                short_name = folder_names[j].split()[0].capitalize()
                            present_folders.append(short_name)
                    
                    print(f"[UPSET DEBUG] Found {len(present_folders)} present folders: {present_folders}")
                    
                    # Calculate the correct label index (bar_idx is now correct since we filtered out the first bar)
                    label_idx = bar_idx
                    
                    # Use label_idx to cycle through ALL folder names
                    if all_folder_names:
                        label = all_folder_names[label_idx % len(all_folder_names)]
                        y_label = len(df_for_upset.columns) - 0.3
                        
                        # Place the label
                        matrix_ax.text(x, y_label, label, ha='center', va='bottom', 
                                      fontsize=10, color='red', rotation=90, 
                                      clip_on=False, weight='bold')
                        
                        print(f"[UPSET DEBUG] Bar (original index: {original_index}, label_idx: {label_idx}): "
                              f"Drawing folder label: '{label}' at x={x:.2f}, "
                              f"intersection: {actual_intersection}, "
                              f"folders in intersection: {present_folders}, "
                              f"assigned label: {label} (from all_folder_names[{label_idx % len(all_folder_names)}])")
                
                # Customize the plot
                plt.title("Concept Groups Overlap Across Folders", fontsize=14, pad=20)
                
                # Save the plot
                upset_plot_path = os.path.join(parent, "groups_upset_plot.png")
                plt.savefig(upset_plot_path, dpi=150, bbox_inches='tight', pad_inches=0.5)
                plt.close()
                
            except Exception as e:
                print(f"[UPSET ERROR] Failed to create UpSet plot: {e}")
                print(f"[UPSET ERROR] Exception type: {type(e)}")
                import traceback
                print(f"[UPSET ERROR] Traceback: {traceback.format_exc()}")                
                # Add overlap summary text to the first page (before the UpSet diagram)

                            # Add the UpSet diagram in a table format (summary on left, diagram on right)
            doc.add_heading("Overlap Summary", level=1)
            plot_table = doc.add_table(rows=1, cols=2)
            plot_table.autofit = False
            
            # Left column for overlap summary text
            left_cell = plot_table.rows[0].cells[0]
            left_cell.width = docx.shared.Inches(3)  # Width for text
            
            # Add overlap summary text to left column
            para = left_cell.paragraphs[0]
            para.add_run("📊 Total concepts: ").bold = True
            run = para.add_run(str(total_concepts))
            run.bold = True
            para.add_run("\n\n")
            
            # Folder information
            for i, name in enumerate(folder_names):
                folder_line = f"📁 {name}: "
                run = para.add_run(folder_line)
                run.bold = True
                run2 = para.add_run(f"{folder_concept_counts[i]} concepts; {folder_group_counts[i]} groups")
                run2.bold = True
                para.add_run("\n")
            
            para.add_run("\n")
            
            # Regrouping method
            para.add_run("🎨 The concepts have been regrouped into ").bold = True
            run = para.add_run(str(num_groups))
            run.bold = True
            para.add_run(" concept groups through the method: ")
            run2 = para.add_run(method_str)
            run2.bold = True
            para.add_run(".\n\n")
            
            # Unique/Shared groups per folder
            for i, folder in enumerate(valid_folders):
                unique_groups = sum(1 for row in color_overlap_table if row[i+1] and sum(row[1:]) == 1)
                shared_groups = sum(1 for row in color_overlap_table if all(row[1:]))
                percent_unique = 100.0 * unique_groups / total_unique_groups if total_unique_groups else 0
                percent_shared = 100.0 * shared_groups / total_unique_groups if total_unique_groups else 0
                
                para.add_run(f"🟢 {folder_names[i]}: ").bold = True
                run = para.add_run(f"{unique_groups}")
                run.bold = True
                para.add_run(" unique concept groups (")
                run = para.add_run(f"{percent_unique:.1f}%")
                run.bold = True
                para.add_run("), ")
                run = para.add_run(f"{shared_groups}")
                run.bold = True
                para.add_run(" shared concept groups (")
                run = para.add_run(f"{percent_shared:.1f}%")
                run.bold = True
                para.add_run(")\n")
            
            # Right column for the plot
            right_cell = plot_table.rows[0].cells[1]
            right_cell.width = docx.shared.Inches(7)  # Width for plot
            right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = right_cell.paragraphs[0].add_run()
            
            # Calculate height to fit page
            page_height_inches = 8.5  # 8.5 inches * 72 points per inch
            margin_inches = 1.0  # 1 inch margins
            available_height_inches = page_height_inches - (2 * margin_inches)
            
            run.add_picture(upset_plot_path, height=Inches(available_height_inches))
                
            # Add Group Presence Matrix on the second page

            doc.add_page_break()
            doc.add_heading("Group Presence Matrix", level=2)
            group_table = doc.add_table(rows=1, cols=1+len(folder_names))
            group_table.rows[0].cells[0].text = "Group"
            for i, folder_name in enumerate(folder_names):
                # Use the same folder names as in the Unique/Common Concepts Table
                group_table.rows[0].cells[1+i].text = folder_name
            
            # Make headings bold and smaller font
            for cell in group_table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.size = docx.shared.Pt(9)  # Smaller font
            
            # Add data rows with colored concepts and smaller font
            for i, row_data in enumerate(group_data):
                row = group_table.add_row().cells
                # Use the unique group name from the DataFrame
                group_name = unique_group_names[i] if i < len(unique_group_names) else row_data[0]
                row[0].text = group_name
                
                # Set smaller font for all cells in this row
                for cell in row:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = docx.shared.Pt(8)  # Even smaller font
                
                # Color the group name based on the group's color
                if self.llm_grouping_var.get():
                    # For LLM grouping, find the color from llm_group_tuples
                    for color, concepts in llm_group_tuples:
                        if extract_group_name(concepts) == group_name:
                            if color.startswith('#') and len(color) == 7:
                                r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                                row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(r, g, b)
                            break
                else:
                    # For color grouping, find the color from color_to_concepts
                    for color, concepts in color_to_concepts.items():
                        if extract_group_name(concepts) == group_name:
                            if color.startswith('#') and len(color) == 7:
                                r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                                row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(r, g, b)
                            break
                
                # Add presence indicators with green/red emoticons
                for j, present in enumerate(row_data[1:]):
                    row[1+j].text = "🟢" if present else "🔴"

            # --- Unique/Common Concepts Table ---
            doc.add_heading("Unique/Common Concepts Table", level=2)
            uniq_table = doc.add_table(rows=1, cols=5)
            widths = [1.5, 4.5, 1.2, 1.5, 2.0]  # inches - adjusted for new Group column
            for i, w in enumerate(widths):
                uniq_table.columns[i].width = docx.shared.Inches(w)
            # Make headings bold
            for cell in uniq_table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
            uniq_table.rows[0].cells[0].text = "Group"
            uniq_table.rows[0].cells[1].text = "Concepts"
            uniq_table.rows[0].cells[2].text = "Unique/Shared"
            uniq_table.rows[0].cells[3].text = "Folders"
            uniq_table.rows[0].cells[4].text = "#Concepts"
            
            folder_unique_words_map = {}
            all_concepts_table = {}
            for color, concepts in (llm_group_tuples if self.llm_grouping_var.get() else color_to_concepts.items()):
                for concept in concepts:
                    folders_with_concept = set()
                    for i, folder in enumerate(valid_folders):
                        if concept in folder_concepts[folder]:
                            folders_with_concept.add(folder_names[i])
                    unique_words = set()
                    for fname in folders_with_concept:
                        unique_words |= folder_unique_map[fname]
                    all_concepts_table[concept] = [w.capitalize() for w in unique_words]
            
            if self.llm_grouping_var.get():
                for idx, (color, concepts) in enumerate(llm_group_tuples):
                    row = uniq_table.add_row().cells
                    
                    # Group name column
                    group_name = extract_group_name(concepts)
                    row[0].text = group_name
                    
                    # Concepts column (colored)
                    para = row[1].paragraphs[0]
                    for cidx, concept in enumerate(concepts):
                        rc = para.add_run(concept)
                        if color.startswith('#') and len(color) == 7:
                            r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                            rc.font.color.rgb = RGBColor(r, g, b)
                        if cidx < len(concepts) - 1:
                            para.add_run(", ")
                    
                    present_folders = [folder_names[i] for i, present in enumerate([row2 for row2 in color_overlap_table if row2[0]==color][0][1:]) if present]
                    n_present = len(present_folders)
                    n_total = len(folder_names)
                    if n_present == 1:
                        row[2].text = "Unique"
                    elif n_present == n_total:
                        row[2].text = f"Common in {n_present} / {n_total} authors"
                    else:
                        row[2].text = f"Partial in {n_present} / {n_total} authors"
                    
                    group_unique_words = set()
                    for concept in concepts:
                        group_unique_words.update(all_concepts_table.get(concept, []))
                    row[3].text = ", ".join(sorted(group_unique_words))
                    row[4].text = str(len(concepts))
            else:
                for idx, color in enumerate(unique_color_groups):
                    row = uniq_table.add_row().cells
                    concepts = color_to_concepts[color]
                    
                    # Group name column
                    group_name = extract_group_name(concepts)
                    row[0].text = group_name
                    
                    # Concepts column (colored)
                    para = row[1].paragraphs[0]
                    for cidx, concept in enumerate(concepts):
                        rc = para.add_run(concept)
                        if color.startswith('#') and len(color) == 7:
                            r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                            rc.font.color.rgb = RGBColor(r, g, b)
                        if cidx < len(concepts) - 1:
                            para.add_run(", ")
                    
                    present_folders = [folder_names[i] for i, present in enumerate([row2 for row2 in color_overlap_table if row2[0]==color][0][1:]) if present]
                    n_present = len(present_folders)
                    n_total = len(folder_names)
                    if n_present == 1:
                        row[2].text = "Unique"
                    elif n_present == n_total:
                        row[2].text = f"Common in {n_present} / {n_total} authors"
                    else:
                        row[2].text = f"Partial in {n_present} / {n_total} authors"
                    
                    group_unique_words = set()
                    for concept in concepts:
                        group_unique_words.update(all_concepts_table.get(concept, []))
                    row[3].text = ", ".join(sorted(group_unique_words))
                    row[4].text = str(len(concepts))


            doc.add_heading("Aggregated Results", level=1)
            # Table of images
            doc.add_heading("UpSet Plots from Each Folder", level=2)
            img_table = doc.add_table(rows=1, cols=5)
            hdr = img_table.rows[0].cells
            hdr[0].text = "Folder"
            hdr[1].text = "BM25"
            hdr[2].text = "Top-k"
            hdr[3].text = "Top-p"
            hdr[4].text = "Temp"
            for folder_name, imgs in folder_images:
                row = img_table.add_row().cells
                row[0].text = folder_name
                for i, img_path in enumerate(imgs):
                    try:
                        run = row[i+1].paragraphs[0].add_run()
                        run.add_picture(img_path, width=Inches(1.5))
                    except Exception as e:
                        row[i+1].text = "[Image error]"
            doc.add_page_break()
            # Table of all canonical concepts (color-coded, with variants, grouped by color)
            doc.add_heading("All Concepts (Grouped by Color, Color-coded)", level=2)
            # color_to_concepts is already created above
            concept_table = doc.add_table(rows=1, cols=3)
            concept_table.rows[0].cells[0].text = "Color Group"
            concept_table.rows[0].cells[1].text = "Concepts"
            concept_table.rows[0].cells[2].text = "Folder Unique Words"
            for color, concepts in color_to_concepts.items():
                row = concept_table.add_row().cells
                run = row[0].paragraphs[0].add_run(color)
                if color.startswith('#') and len(color) == 7:
                    r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                    run.font.color.rgb = RGBColor(r, g, b)
                # All concepts in this group in the same color
                para = row[1].paragraphs[0]
                for idx, c in enumerate(concepts):
                    rc = para.add_run(c)
                    color_val = group_colors.get(c, '#FF0000')
                    if c not in group_colors:
                        print(f"[GROUP COLOR WARNING] Concept not in group_colors: {c}")
                        self.root.after(0, lambda c=c: self.aggregate_status_label.config(text=f"[GROUP COLOR WARNING] Concept not in group_colors: {c}", foreground='red'))
                        group_colors[c] = '#FF0000'
                    if color_val.startswith('#') and len(color_val) == 7:
                        rc.font.color.rgb = RGBColor(int(color_val[1:3], 16), int(color_val[3:5], 16), int(color_val[5:7], 16))
                    if idx < len(concepts) - 1:
                        para.add_run(", ")
                # Folder unique words for this group (use same logic as above)
                folders_with_concept = set()
                for c in concepts:
                    for i, folder in enumerate(valid_folders):
                        if c in folder_concepts[folder]:
                            folders_with_concept.add(folder_names[i])
                unique_words = set()
                for fname in folders_with_concept:
                    unique_words |= folder_unique_map[fname]
                row[2].text = ", ".join(sorted(unique_words))
            doc.add_page_break()
            # --- Overlap table: put all concepts in the same color group on the same row ---
            doc.add_heading("Concept Overlap Across Folders (Grouped)", level=2)
            overlap_doc_table = doc.add_table(rows=1, cols=1+len(folder_names))
            hdr = overlap_doc_table.rows[0].cells
            hdr[0].text = "Concepts in Group"
            for i, name in enumerate(folder_names):
                hdr[1+i].text = name
            # Make headings bold
            for cell in overlap_doc_table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
            if self.llm_grouping_var.get():
                for idx, (color, concepts) in enumerate(llm_group_tuples):
                    doc_row = overlap_doc_table.add_row().cells
                    para = doc_row[0].paragraphs[0]
                    for cidx, c in enumerate(concepts):
                        rc = para.add_run(c)
                        if color.startswith('#') and len(color) == 7:
                            r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                            rc.font.color.rgb = RGBColor(r, g, b)
                        if cidx < len(concepts) - 1:
                            para.add_run(", ")
                    concepts_in_color = set(concepts)
                    for j, folder in enumerate(valid_folders):
                        present = any(c in folder_concepts[folder] for c in concepts_in_color)
                        doc_row[1+j].text = "✔" if present else ""
            else:
                for idx, color in enumerate(unique_color_groups):
                    doc_row = overlap_doc_table.add_row().cells
                    para = doc_row[0].paragraphs[0]
                    concepts = color_to_concepts[color]
                    for cidx, c in enumerate(concepts):
                        rc = para.add_run(c)
                        if color.startswith('#') and len(color) == 7:
                            r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                            rc.font.color.rgb = RGBColor(r, g, b)
                        if cidx < len(concepts) - 1:
                            para.add_run(", ")
                    concepts_in_color = set(concepts)
                    for j, folder in enumerate(valid_folders):
                        present = any(c in folder_concepts[folder] for c in concepts_in_color)
                        doc_row[1+j].text = "✔" if present else ""
            safe_update_progress(100, total_elapsed, total_elapsed)
            # Save
            out_path = os.path.join(parent, "aggregated_results.docx")
            doc.save(out_path)
            safe_update_status(f"Aggregation complete. Saved to {out_path}", "green")
        except Exception as e:
            tb = traceback.format_exc()
            safe_update_status(f"Error generating DOCX: {e}", "red")
            print(tb)

    def _update_aggregate_progress(self, val, elapsed, est_total):
        self.aggregate_progress['value'] = val
        self.aggregate_time_label.config(text=f"Elapsed: {elapsed:.1f}s, Estimated total: {est_total:.1f}s")

    def on_llm_grouping_toggle(self):
        # If LLM Grouping is enabled, disable all other grouping controls
        state = 'disabled' if self.llm_grouping_var.get() else 'normal'
        self.agg_color_criteria_cb.config(state=state)
        self.agg_min_letters_entry.config(state=state)
        self.agg_group_by_subletters.set(False if self.llm_grouping_var.get() else self.agg_group_by_subletters.get())
        self.agg_group_by_words.set(False if self.llm_grouping_var.get() else self.agg_group_by_words.get())
        self.sim_threshold_entry.config(state=state)
        self.grouping_logic_combo.config(state=state)
        self.agg_group_by_subletters_cb_state = getattr(self, 'agg_group_by_subletters_cb_state', None)
        self.agg_group_by_words_cb_state = getattr(self, 'agg_group_by_words_cb_state', None)
        # Optionally, disable the checkboxes themselves if you want
        # (requires storing references to the checkboxes)
        # self.agg_group_by_subletters_cb.config(state=state)
        # self.agg_group_by_words_cb.config(state=state)

def replace_bekker_with_greek(concept, bekker_map):
    for bekker, greek in bekker_map.items():
        if bekker in concept:
            # Replace the Bekker number with the Greek concept
            return concept.replace(bekker, greek)
    return concept

def convert_docx_to_pdf_libreoffice(docx_path, output_dir=None):
    if output_dir is None:
        output_dir = os.path.dirname(docx_path)
    try:
        result = subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            docx_path,
            "--outdir", output_dir
        ], check=True, capture_output=True, text=True)
        print("LibreOffice output:", result.stdout)
        pdf_path = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF not found at {pdf_path}")
        return pdf_path
    except Exception as e:
        print("Error converting DOCX to PDF with LibreOffice:", e)
        return None

def extract_groups_from_llm_output(parsed, concepts):
    # If it's a dict, use values
    if isinstance(parsed, dict):
        return list(parsed.values())
    # If it's a list of dicts with 'concepts' key, extract those
    if isinstance(parsed, list) and all(isinstance(g, dict) and 'concepts' in g for g in parsed):
        return [g['concepts'] for g in parsed]
    # If it's a list of lists, use as-is
    if isinstance(parsed, list) and all(isinstance(g, list) for g in parsed):
        return parsed
    # If it's a flat list, treat each as a singleton group
    if isinstance(parsed, list):
        return [[g] for g in parsed]
    # Fallback: treat all as one group
    return [concepts]

def real_llm_grouping(concepts, prompt, model, output_dir=None):
    """Call the real LLM API for grouping. Returns the LLM's raw output and parsed groups."""
    # Prepare the prompt
    prompt_full = (
        prompt.strip() +
        "\nIf the output is too large, split it into multiple JSON arrays or use short group names. Return only JSON, no explanation.\nConcepts:\n" +
        "\n".join(f"- {c}" for c in concepts)
    )
    print("[LLM GROUPING REAL] Model:", model)
    print("[LLM GROUPING REAL] Prompt:\n", prompt_full)
    print("[LLM GROUPING REAL] Concepts:", concepts)
    llm_output = None
    parsed_groups = None
    error = None
    input_tokens = len(prompt_full.split())
    output_tokens = 0
    all_extracted_groups = []
    try:
        # --- Model selection and mapping logic (from advanced_rag.py) ---
        model_map = {
            "gpt-3.5": "gpt-3.5-turbo",
            "gpt-4o": "gpt-4o",
            "gpt-4o mini": "gpt-4o-mini",
            "o1-mini": "gpt-4o-mini",
            "o3-mini": "gpt-4o-mini",
            "mistral": "mistral-small-latest",
            "mistral-api": "mistral-small-latest",
            "meta-llama-3": "meta-llama/Meta-Llama-3-8B-Instruct",
            "remote meta-llama-3": "meta-llama/Meta-Llama-3-8B-Instruct",
            "qwen3": "Qwen/Qwen1.5-7B-Chat",
            # Add more mappings as needed
        }
        # Max tokens per model (based on public docs)
        max_tokens_map = {
            "gpt-3.5": 16385,
            "gpt-4o": 128000,
            "gpt-4o mini": 128000,
            "o1-mini": 128000,
            "o3-mini": 128000,
            "mistral": 128000,
            "mistral-api": 128000,
            "meta-llama-3": 128000,
            "remote meta-llama-3": 128000,
            "qwen3": 128000,
            "gemini": 128000,
            "claude": 200000,
        }
        normalized = model.lower().replace("-api", "").replace("remote ", "").replace(" ", "-")
        model_key = None
        for key in model_map:
            if key in normalized:
                model_key = key
                break
        # Default max tokens
        max_tokens = max_tokens_map.get(model_key, 128000)
        # --- OpenAI GPT Models ---
        if "gpt" in model.lower() or "o1-mini" in model.lower() or "o3-mini" in model.lower():
            # Use 128000 for GPT-4o and variants
            openai_api_key = os.environ.get("OPENAI_API_KEY")
            if not openai_api_key:
                raise RuntimeError("OpenAI API key not available.")
            import openai
            client = openai.OpenAI(api_key=openai_api_key)
            model_name = model_map.get(model_key, "gpt-3.5-turbo")
            response = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt_full}],
                temperature=0.3,
                max_tokens=max_tokens  # 128000 for GPT-4o, 16385 for GPT-3.5
            )
            llm_output = response.choices[0].message.content
            try:
                import tiktoken
                enc = tiktoken.encoding_for_model(model_name)
                input_tokens = len(enc.encode(prompt_full))
                output_tokens = len(enc.encode(llm_output))
            except Exception:
                input_tokens = len(prompt_full.split())
                output_tokens = len(llm_output.split())
        elif "mistral" in model.lower():
            # Use 128000 for Mistral
            api_key = os.environ.get("MISTRAL_API_KEY")
            if not api_key or not Mistral:
                raise RuntimeError("Mistral API or library not available.")
            client = Mistral(api_key=api_key)
            response = client.chat.complete(
                model=model_map.get(model_key, "mistral-small-latest"),
                messages=[{"role": "user", "content": prompt_full}],
                temperature=0.3,
                max_tokens=max_tokens  # 128000
            )
            llm_output = response.choices[0].message.content
            input_tokens = len(prompt_full.split())
            output_tokens = len(llm_output.split())
        elif "llama" in model.lower():
            # Use 128000 for Llama
            hf_token = os.environ.get("HF_API_TOKEN")
            if not hf_token or not InferenceClient:
                raise RuntimeError("HuggingFace Inference API or library not available.")
            client = InferenceClient(token=hf_token, timeout=120)
            llm_output = client.text_generation(
                prompt_full,
                model=model_map.get(model_key, "meta-llama/Meta-Llama-3-8B-Instruct"),
                temperature=0.3,
                max_new_tokens=max_tokens  # 128000
            )
            input_tokens = len(prompt_full.split())
            output_tokens = len(llm_output.split())
        elif "qwen3" in model.lower():
            # Use 128000 for Qwen3
            try:
                from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline
            except ImportError:
                raise RuntimeError("transformers library not installed. Please install with 'pip install transformers'.")
            model_id = model_map.get(model_key, "Qwen/Qwen1.5-7B-Chat")
            hf_token = os.environ.get("HF_API_TOKEN")
            print(f"[LLM GROUPING REAL] Loading Qwen3 model {model_id} (this may take a while if not cached)...")
            try:
                tokenizer = AutoTokenizer.from_pretrained(model_id, token=hf_token)
                model_qwen = AutoModelForCausalLM.from_pretrained(model_id, token=hf_token)
            except Exception as e:
                raise RuntimeError(f"Qwen3 model not found or access denied: {e}")
            pipe = pipeline("text-generation", model=model_qwen, tokenizer=tokenizer)
            result = pipe(prompt_full, max_new_tokens=max_tokens, do_sample=True, temperature=0.3)  # 128000
            llm_output = result[0]["generated_text"]
            input_tokens = len(prompt_full.split())
            output_tokens = len(llm_output.split())
        elif "gemini" in model.lower():
            # Use 128000 for Gemini
            try:
                import google.generativeai as genai
            except ImportError:
                print("[LLM GROUPING REAL] google-generativeai not installed.")
                raise RuntimeError("google-generativeai not installed.")
            gemini_api_key = os.environ.get("GEMINI_API_KEY")
            if not gemini_api_key:
                raise RuntimeError("GEMINI_API_KEY not set in environment.")
            genai.configure(api_key=gemini_api_key)
            model_name = "models/gemini-1.5-pro-latest"
            try:
                model_gemini = genai.GenerativeModel(model_name)
                response = model_gemini.generate_content(prompt_full, generation_config={"max_output_tokens": max_tokens})  # 128000
                llm_output = response.text
            except Exception as e:
                raise RuntimeError(f"Gemini API error: {e}")
            input_tokens = len(prompt_full.split())
            output_tokens = len(llm_output.split())
        elif "claude" in model.lower():
            # Use 200000 for Claude
            try:
                import anthropic
            except ImportError:
                print("[LLM GROUPING REAL] anthropic not installed.")
                raise RuntimeError("anthropic not installed.")
            claude_api_key = os.environ.get("ANTHROPIC_API_KEY")
            if not claude_api_key:
                raise RuntimeError("ANTHROPIC_API_KEY not set in environment.")
            client = anthropic.Anthropic(api_key=claude_api_key)
            tried_models = []
            for model_name in [
                "claude-3-7-sonnet-20250224",  # Claude 3.7 Sonnet (Feb 2025)
                "claude-4-sonnet-20250501",    # Claude 4 Sonnet (May 2025)
                "claude-3-5-sonnet-20241022",  # Claude 3.5 Sonnet (Oct 2024)
                "claude-3-haiku-20240307"      # Claude 3 Haiku (Mar 2024)
            ]:
                tried_models.append(model_name)
                try:
                    response = client.messages.create(
                        model=model_name,
                        max_tokens=200000,  # Claude max
                        temperature=0.3,
                        messages=[{"role": "user", "content": prompt_full}]
                    )
                    llm_output = response.content[0].text if hasattr(response.content[0], 'text') else str(response.content)
                    break
                except Exception as e:
                    print(f"[LLM GROUPING REAL] Claude model {model_name} not available: {e}")
            else:
                raise RuntimeError(f"Claude API error: None of the tried models are available: {tried_models}")
            input_tokens = len(prompt_full.split())
            output_tokens = len(llm_output.split())
        elif "grok" in model.lower():
            print("[LLM GROUPING REAL] Grok 3 and Grok 4 models are not available via public API.")
            raise RuntimeError("Grok 3 and Grok 4 models are not yet implemented (no public API).")
        elif "phi4" in model.lower():
            # --- Phi4 (Not available) ---
            print("[LLM GROUPING REAL] Phi4 model not yet implemented (no public API).")
            raise RuntimeError("Phi4 model not yet implemented (no public API).")
        elif "meta llama 70b" in model.lower():
            # --- Meta Llama 70B (Not available) ---
            print("[LLM GROUPING REAL] Meta Llama 70B model not yet implemented (no public API).")
            raise RuntimeError("Meta Llama 70B model not yet implemented (no public API).")
        elif "deepseek" in model.lower():
            # --- DeepSeek V3 (Not available) ---
            print("[LLM GROUPING REAL] DeepSeek V3 model not yet implemented (no public API).")
            raise RuntimeError("DeepSeek V3 model not yet implemented (no public API).")
        elif "nebius" in model.lower():
            # --- Mistral (Nebius) (Not available) ---
            print("[LLM GROUPING REAL] Mistral (Nebius) model not yet implemented (no public API).")
            raise RuntimeError("Mistral (Nebius) model not yet implemented (no public API).")
        else:
            raise RuntimeError(f"Unsupported model: {model}")
        print("[LLM GROUPING REAL] LLM Output:\n", llm_output)
        print(f"[LLM GROUPING REAL] Input tokens: {input_tokens}, Output tokens: {output_tokens}")
        # Try to parse JSON from the output
        import re
        # Extract all JSON arrays or objects from the output
        json_matches = re.findall(r'\{[\s\S]*?\}|\[[\s\S]*?\]', llm_output)
        for match in json_matches:
            try:
                obj = json.loads(match)
                # If it's a dict of groups, convert to list of lists
                if isinstance(obj, dict):
                    group_list = list(obj.values())
                    all_extracted_groups.extend(group_list)
                elif isinstance(obj, list):
                    all_extracted_groups.extend(obj)
            except Exception as e:
                print("[LLM GROUPING REAL] JSON parse error in match:", e)
        if all_extracted_groups:
            parsed_groups = extract_groups_from_llm_output(all_extracted_groups, concepts)
        else:
            # Try to parse the whole output as JSON
            try:
                parsed_groups = json.loads(llm_output)
                parsed_groups = extract_groups_from_llm_output(parsed_groups, concepts)
            except Exception as e:
                print("[LLM GROUPING REAL] Fallback JSON parse error:", e)
                parsed_groups = [concepts]
        # Warn if output is likely truncated
        if len(llm_output) > 18000 or (llm_output and not llm_output.rstrip().endswith(']') and not llm_output.rstrip().endswith('}')):
            print("[LLM GROUPING REAL] WARNING: Output may be truncated or incomplete!")
        # After parsing and flattening, print each group and its concepts
        print("[LLM GROUPING REAL] Parsed groups:")
        for idx, group in enumerate(parsed_groups):
            print(f"  Group {idx+1}: {group}")
    except Exception as e:
        error = str(e)
        print("[LLM GROUPING REAL] ERROR:", error)
        parsed_groups = [concepts]
        llm_output = error
    # Save output to file
    if output_dir:
        try:
            with open(os.path.join(output_dir, "llm_grouping_output.json"), "w", encoding="utf-8") as f:
                json.dump(parsed_groups, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print("[LLM GROUPING REAL] Failed to save output:", e)
    print(f"[LLM GROUPING REAL] Parsed {len(parsed_groups)} groups.")
    return parsed_groups

def main():
    root = tk.Tk()
    app = UpSetGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()