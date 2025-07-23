import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import pandas as pd
from sklearn.preprocessing import MultiLabelBinarizer
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from upsetplot import UpSet, from_indicators
import numpy as np
import os
import warnings
from PIL import Image, ImageTk, ImageDraw, ImageFont
import threading
from collections import defaultdict
import re
import html
import traceback
import textwrap
import win32com.client
import subprocess

# Common suffixes for substring grouping (except suffixes)
common_suffixes = [
    'ation', 'ment', 'ness', 'sion', 'tion', 'ing', 'ed', 'ly', 'er', 'est', 'ful', 'less', 'able', 'ible', 'ous', 'ive', 'al', 'ic', 'ant', 'ent', 'ism', 'ist', 'ity', 'ty', 'en', 'ize', 'ise', 'ward', 'wise'
]

warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", category=FutureWarning)

def wrap_label(text, width=18):
    # Try to wrap at word boundaries, fallback to hard wrap
    return '\n'.join(textwrap.wrap(text, width=width))

class ZoomablePanCanvas:
    def __init__(self, parent):
        self.parent = parent
        self.canvas = tk.Canvas(parent, bg='white')
        
        # Scrollbars
        self.h_scrollbar = ttk.Scrollbar(parent, orient="horizontal", command=self.canvas.xview)
        self.v_scrollbar = ttk.Scrollbar(parent, orient="vertical", command=self.canvas.yview)
        self.canvas.configure(xscrollcommand=self.h_scrollbar.set, yscrollcommand=self.v_scrollbar.set)
        
        # Grid layout
        self.canvas.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.h_scrollbar.grid(row=1, column=0, sticky=(tk.W, tk.E))
        self.v_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # Configure grid weights
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)
        
        # Zoom and pan variables
        self.zoom_level = 1.0
        self.original_images = []
        self.image_refs = []
        self.image_items = []
        
        # Bind mouse events for panning
        self.canvas.bind("<Button-1>", self.start_pan)
        self.canvas.bind("<B1-Motion>", self.do_pan)
        self.canvas.bind("<MouseWheel>", self.zoom_wheel)
        
        self.last_x = 0
        self.last_y = 0
        
    def start_pan(self, event):
        self.canvas.scan_mark(event.x, event.y)
        self.last_x = event.x
        self.last_y = event.y
        
    def do_pan(self, event):
        self.canvas.scan_dragto(event.x, event.y, gain=1)
        
    def zoom_wheel(self, event):
        if event.delta > 0:
            self.zoom_in()
        else:
            self.zoom_out()
            
    def set_images(self, image_paths, window_width):
        self.original_images = []
        self.image_refs = []
        self.image_items = []
        self.canvas.delete("all")
        if not image_paths:
            return
        loaded_imgs = []
        max_height = 0
        for img_path in image_paths:
            try:
                img = Image.open(img_path)
                loaded_imgs.append(img)
                max_height = max(max_height, img.height)
            except Exception as e:
                print(f"Error loading image {img_path}: {e}")
        x_offset = 10
        padding = 10
        for img in loaded_imgs:
            # Align to bottom
            y_offset = max_height - img.height + padding
            self.original_images.append((img, x_offset, y_offset))
            x_offset += img.width + padding
        self.zoom_level = 1.0
        self.update_display()
        
    def update_display(self):
        self.canvas.delete("all")
        self.image_refs = []
        self.image_items = []
        
        for img, x, y in self.original_images:
            # Apply zoom
            new_width = int(img.width * self.zoom_level)
            new_height = int(img.height * self.zoom_level)
            
            if new_width > 0 and new_height > 0:
                zoomed_img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                photo = ImageTk.PhotoImage(zoomed_img)
                self.image_refs.append(photo)
                
                # Position with zoom
                scaled_x = int(x * self.zoom_level)
                scaled_y = int(y * self.zoom_level)
                
                item = self.canvas.create_image(scaled_x, scaled_y, anchor="nw", image=photo)
                self.image_items.append(item)
        
        # Update scroll region
        self.canvas.update_idletasks()
        bbox = self.canvas.bbox("all")
        if bbox:
            self.canvas.configure(scrollregion=bbox)
            
    def zoom_in(self):
        self.zoom_level = min(self.zoom_level * 1.1, 5.0)
        self.update_display()
        
    def zoom_out(self):
        self.zoom_level = max(self.zoom_level / 1.1, 0.2)
        self.update_display()
        
    def fit_to_window(self):
        if not self.original_images:
            return
        # Calculate total width of all images + padding
        total_width = sum(img.width for img, _, _ in self.original_images)
        total_width += 10 * (len(self.original_images) + 1)
        canvas_width = self.canvas.winfo_width()
        if canvas_width <= 1:  # Not yet rendered
            self.parent.update_idletasks()
            canvas_width = self.canvas.winfo_width()
        if total_width > 0 and canvas_width > 0:
            self.zoom_level = min(1.0, canvas_width / total_width)
        else:
            self.zoom_level = 1.0
        self.update_display()

class UpSetGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("UpSet Plot Generator")
        self.root.geometry("1400x800")
        
        self.csv_file = None
        self.output_images = []
        self.color_mapping = {}
        self.use_colors = tk.BooleanVar(value=True)
        self.group_by_subletters = tk.BooleanVar(value=False)
        self.group_by_words = tk.BooleanVar(value=True)
        self.group_by_same_color = tk.BooleanVar(value=False)
        
        self.setup_ui()
        
    def setup_ui(self):
        # Main frame
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.columnconfigure(2, weight=1)
        main_frame.rowconfigure(5, weight=1)  # Make results frame expand
        
        # --- Merge CSVs section ---
        merge_frame = ttk.LabelFrame(main_frame, text="Merge CSV Files", padding="5")
        merge_frame.grid(row=0, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0,2))
        self.merge_label = ttk.Label(merge_frame, text="No files selected", foreground="gray")
        self.merge_label.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5), pady=2)
        # Place the button below the label, natural width
        self.merge_button = ttk.Button(merge_frame, text="Select & Merge CSVs", command=self.merge_csv_files)
        self.merge_button.grid(row=1, column=0, sticky=tk.W, padx=(0, 5), pady=2)
        
        # File selection
        ttk.Label(main_frame, text="CSV File:").grid(row=1, column=0, sticky=tk.W, pady=0)
        self.file_label = ttk.Label(main_frame, text="No file selected", foreground="gray")
        self.file_label.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=(5, 0), pady=0)
        ttk.Button(main_frame, text="Browse", command=self.browse_file).grid(row=1, column=2, padx=(5, 0), pady=0)
        # Block info label (single line, under file label)
        self.block_info_label = ttk.Label(main_frame, text="", foreground="blue")
        self.block_info_label.grid(row=2, column=1, sticky=(tk.W, tk.E), padx=(5, 0), pady=0)
        
        # Options frame
        options_frame = ttk.LabelFrame(main_frame, text="Options", padding="3")
        options_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=0)
        
        # Color checkbox and entry
        color_frame = ttk.Frame(options_frame)
        color_frame.grid(row=0, column=0, sticky=tk.W, pady=0)
        self.color_criteria_cb = ttk.Checkbutton(color_frame, text="Color criteria", variable=self.use_colors, command=self.on_color_criteria_toggle)
        self.color_criteria_cb.grid(row=0, column=0, sticky=tk.W, pady=0)
        ttk.Label(color_frame, text="Min shared letters:").grid(row=0, column=1, padx=(5,0), pady=0)
        self.min_letters_var = tk.StringVar(value="5")
        self.min_letters_entry = ttk.Entry(color_frame, textvariable=self.min_letters_var, width=3)
        self.min_letters_entry.grid(row=0, column=2, padx=(2,0), pady=0)
        ttk.Checkbutton(color_frame, text="Group by subletters (except suffixes)", variable=self.group_by_subletters, command=self.on_group_by_subletters).grid(row=0, column=3, padx=(5,0), pady=0)
        ttk.Checkbutton(color_frame, text="Group by whole words", variable=self.group_by_words, command=self.on_group_by_words).grid(row=0, column=4, padx=(5,0), pady=0)
        self.group_by_same_color_cb = ttk.Checkbutton(color_frame, text="Collapse into same color group", variable=self.group_by_same_color)
        self.group_by_same_color_cb.grid(row=0, column=5, padx=(5,0), pady=0)
        self.group_by_same_color_cb.state(['disabled'])
        # Add label wrap width entry
        ttk.Label(color_frame, text="Label wrap width:").grid(row=0, column=6, padx=(5,0), pady=0)
        self.wrap_width_var = tk.StringVar(value="75")
        self.wrap_width_entry = ttk.Entry(color_frame, textvariable=self.wrap_width_var, width=3)
        self.wrap_width_entry.grid(row=0, column=7, padx=(2,0), pady=0)
        
        # Process button and progress
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=0)
        button_frame.columnconfigure(0, weight=1)
        
        self.progress = ttk.Progressbar(button_frame, mode='indeterminate')
        self.progress.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5), pady=0)
        
        self.process_btn = ttk.Button(button_frame, text="Generate UpSet Plots", 
                                     command=self.start_processing, state='disabled')
        self.process_btn.grid(row=0, column=1, pady=0)
        
        # Results frame with zoom controls
        results_frame = ttk.LabelFrame(main_frame, text="Results", padding="5")
        results_frame.grid(row=5, column=0, columnspan=3, sticky=(tk.N, tk.S, tk.E, tk.W), pady=5)
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(1, weight=1)
        
        # Zoom controls
        zoom_frame = ttk.Frame(results_frame)
        zoom_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 2))
        
        ttk.Button(zoom_frame, text="Zoom In", command=self.zoom_in).grid(row=0, column=0, padx=(0, 2), pady=0)
        ttk.Button(zoom_frame, text="Zoom Out", command=self.zoom_out).grid(row=0, column=1, padx=(0, 2), pady=0)
        ttk.Button(zoom_frame, text="Fit to Window", command=self.fit_to_window).grid(row=0, column=2, padx=(0, 2), pady=0)
        
        ttk.Label(zoom_frame, text="Use mouse wheel to zoom, click and drag to pan").grid(row=0, column=3, padx=(10, 0), pady=0)
        
        # Zoomable canvas
        canvas_frame = ttk.Frame(results_frame)
        canvas_frame.grid(row=1, column=0, sticky=(tk.N, tk.S, tk.E, tk.W))
        canvas_frame.columnconfigure(0, weight=1)
        canvas_frame.rowconfigure(0, weight=1)
        
        self.zoomable_canvas = ZoomablePanCanvas(canvas_frame)
        self.zoomable_canvas.canvas.grid(sticky=(tk.N, tk.S, tk.E, tk.W))  # Ensure canvas expands
        
        # Status label
        self.status_label = ttk.Label(main_frame, text="Ready")
        self.status_label.grid(row=6, column=0, columnspan=3, pady=2)
        
    def browse_file(self):
        file_path = filedialog.askopenfilename(
            title="Select CSV file",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        
        if file_path:
            self.csv_file = file_path
            self.file_label.config(text=os.path.basename(file_path), foreground="black")
            self.process_btn.config(state='normal')
            # Try to read and display block info
            self.display_block_info()
            
    def display_block_info(self):
        # Try to read the CSV and infer block structure
        if not self.csv_file:
            self.block_info_label.config(text="", foreground="blue")
            return
        encodings_to_try = ['cp1253', 'utf-8', 'windows-1252']
        df = None
        for enc in encodings_to_try:
            try:
                df = pd.read_csv(self.csv_file, encoding=enc)
                break
            except Exception:
                continue
        if df is None:
            self.block_info_label.config(text="Could not read CSV to infer blocks.", foreground="red")
            return
        # Detect block structure
        param_names = [col for col in ['Temperature', 'Top-p', 'Top-k', 'BM25 Weight'] if col in df.columns]
        block_size = None
        blocks = []
        if len(param_names) == 4:
            n = len(df)
            for p in param_names:
                values = df[p].values
                first_val = values[0]
                block_len = 1
                for v in values[1:]:
                    if v == first_val:
                        block_len += 1
                    else:
                        break
                if n % block_len == 0 and block_len > 1:
                    block_size = block_len
                    break
            if block_size is None:
                block_info = "Could not infer block size."
            else:
                num_blocks = n // block_size
                blocks = []
                block_strs = []
                for i in range(num_blocks):
                    start = i * block_size
                    end = (i + 1) * block_size
                    varying_param = None
                    for p in param_names:
                        if len(df[p].iloc[start:end].unique()) > 1:
                            varying_param = p
                            break
                    blocks.append((start, end, varying_param if varying_param else '?'))
                    block_strs.append(f"{start}-{end-1}: {varying_param if varying_param else '?'}")
                block_info = "Blocks: [" + ", ".join(block_strs) + "]"
        else:
            block_info = "Could not detect all required parameters."
        self.block_info_label.config(text=block_info, foreground="blue")
        self.detected_blocks = blocks
            
    def zoom_in(self):
        self.zoomable_canvas.zoom_in()
        
    def zoom_out(self):
        self.zoomable_canvas.zoom_out()
        
    def fit_to_window(self):
        self.zoomable_canvas.fit_to_window()
        
    def extract_concepts(self, text):
        if pd.isna(text):
            return []
        lines = text.split('\n')
        concepts = []
        for line in lines:
            line = line.strip()
            if line.startswith(tuple(str(i)+'.' for i in range(1, 21))):
                # Markdown or plain numbered list
                if '**' in line:
                    c = line.split('**')[1].replace(':', '').replace('.', '').strip()
                    if c:
                        concepts.append(c)
                else:
                    c = line.lstrip('0123456789. ').replace(':', '').replace('.', '').strip()
                    if c:
                        concepts.append(c)
        return concepts
    
    def normalize_word(self, word):
        """Normalize word by removing plurals and common variations"""
        word = word.lower().strip()
        
        # Handle common plural forms
        if word.endswith('ies') and len(word) > 4:
            return word[:-3] + 'y'
        elif word.endswith('es') and len(word) > 3:
            return word[:-2]
        elif word.endswith('s') and len(word) > 2:
            return word[:-1]
        
        # Handle common suffixes
        suffixes = ['ing', 'ed', 'er', 'est', 'ly', 'tion', 'sion', 'ness', 'ment', 'able', 'ible', 'ful', 'less']
        for suffix in suffixes:
            if word.endswith(suffix) and len(word) > len(suffix) + 2:
                return word[:-len(suffix)]
        
        return word
    
    def extract_key_words(self, concept):
        """Extract meaningful words from a concept"""
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'should', 'could', 'can', 'may', 'might', 'must'}
        
        words = re.findall(r'\b\w+\b', concept.lower())
        meaningful_words = [self.normalize_word(word) for word in words if word not in stop_words and len(word) > 2]
        
        return meaningful_words
    
    def create_color_mapping(self, all_concepts):
        self.color_mapping = {}
        min_letters = 5
        try:
            min_letters = int(self.min_letters_var.get())
        except Exception:
            pass
        all_concepts_sorted = sorted(all_concepts, key=lambda x: x.lower())
        groups = []
        assigned = set()
        group_by_words = self.group_by_words.get()
        for i, concept in enumerate(all_concepts_sorted):
            concept_lc = concept.lower()
            if concept_lc in assigned:
                continue
            group = [concept]
            words1 = set(re.findall(r'\b\w+\b', concept_lc))
            for j, other in enumerate(all_concepts_sorted):
                other_lc = other.lower()
                if i == j or other_lc in assigned:
                    continue
                words2 = set(re.findall(r'\b\w+\b', other_lc))
                found = False
                if group_by_words:
                    if words1 & words2:
                        found = True
                else:
                    for w1 in words1:
                        for w2 in words2:
                            for k in range(len(w1) - min_letters + 1):
                                sub = w1[k:k+min_letters]
                                if sub and sub in w2:
                                    if (w1.endswith(sub) or w2.endswith(sub)) and sub in common_suffixes:
                                        continue
                                    found = True
                                    break
                            if found:
                                break
                        if found:
                            break
                if found:
                    group.append(other)
            groups.append(group)
            assigned.update([c.lower() for c in group])
        colors = [
            '#800000', '#FF8C00', '#228B22', '#8B008B', '#A0522D', '#2E8B57', '#9932CC', '#FFD700',
            '#556B2F', '#C71585', '#8B4513', '#20B2AA', '#B22222', '#FF4500', '#6A5ACD', '#D2691E',
            '#006400', '#708090', '#FF6347', '#483D8B', '#000000', '#808000', '#8B0000', '#FF1493',
        ]
        color_idx = 0
        for group in groups:
            color = colors[color_idx % len(colors)]
            for concept in group:
                self.color_mapping[concept] = color
            color_idx += 1
        return self.color_mapping
    
    def create_colored_text_image(self, text, color, font_size=12):
        """Create an image with colored text on white background"""
        try:
            font = ImageFont.truetype("arial.ttf", font_size)
        except:
            font = ImageFont.load_default()
        
        # Get text dimensions
        bbox = font.getbbox(text)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        # Create image with white background
        img = Image.new('RGB', (text_width + 10, text_height + 6), 'white')
        draw = ImageDraw.Draw(img)
        
        # Draw text in color
        draw.text((5, 3), text, fill=color, font=font)
        
        return img
    
    def start_processing(self):
        if not self.csv_file:
            messagebox.showerror("Error", "Please select a CSV file first")
            return
            
        # Start processing in a separate thread
        self.process_btn.config(state='disabled')
        self.progress.start()
        self.status_label.config(text="Processing...")
        
        thread = threading.Thread(target=self.process_data)
        thread.daemon = True
        thread.start()
        
    def process_data(self):
        try:
            print('Starting process_data')
            encodings_to_try = ['utf-8', 'windows-1252', 'cp1253', 'latin1']
            for enc in encodings_to_try:
                try:
                    df = pd.read_csv(self.csv_file, encoding=enc)
                    print(f'CSV loaded with encoding: {enc}')
                    break
                except Exception as e:
                    print(f'Failed to load CSV with encoding {enc}: {e}')
            else:
                raise Exception('Could not read CSV file with any supported encoding (utf-8, cp1253, windows-1252, latin1)')
            print('First 10 Main Answer values:')
            print(df['Main Answer'].head(10).to_list())
            df['Concepts'] = df['Main Answer'].apply(self.extract_concepts)
            print('First 10 Concepts values:')
            print(df['Concepts'].head(10).to_list())
            all_concepts = set()
            for concepts in df['Concepts']:
                all_concepts.update(concepts)
            print(f'All unique concepts: {all_concepts}')
            self.create_color_mapping(list(all_concepts))
            print(f'Color mapping: {self.color_mapping}')
            params = ['Temperature', 'Top-p', 'Top-k', 'BM25 Weight']
            param_labels = {'Temperature': 'Temp', 'Top-p': 'Topp', 'Top-k': 'Topk', 'BM25 Weight': 'BM25'}
            outdir = 'compare_gui_output'
            os.makedirs(outdir, exist_ok=True)
            blocks = getattr(self, 'detected_blocks', None)
            if not blocks or len(blocks) == 0:
                n = len(df)
                block_size = n // 4 if n % 4 == 0 else 5
                blocks = [(i*block_size, (i+1)*block_size, params[i] if i < len(params) else '?') for i in range(4)]
            output_files = []
            for start, end, varying_param in blocks:
                print(f'Processing block: {varying_param} ({start}-{end})')
                subset = df.iloc[start:end].copy()
                if subset.empty:
                    print('Subset empty, skipping')
                    continue
                mlb = MultiLabelBinarizer()
                concept_matrix = pd.DataFrame(mlb.fit_transform(subset['Concepts']), columns=mlb.classes_)
                print(f'Concept matrix columns: {concept_matrix.columns}')
                for p in params:
                    concept_matrix[p] = subset[p].values
                concept_matrix_reset = concept_matrix.drop(params, axis=1).astype(bool).reset_index(drop=True)
                print(f'Concept matrix reset columns: {concept_matrix_reset.columns}')
                if self.use_colors.get() and self.group_by_same_color.get():
                    color_map = {col: self.color_mapping.get(col, None) for col in concept_matrix_reset.columns}
                    print(f'Color map: {color_map}')
                    color_groups = defaultdict(list)
                    for col, color in color_map.items():
                        color_groups[color].append(col)
                    print(f'Color groups: {color_groups}')
                    merged = pd.DataFrame(index=concept_matrix_reset.index)
                    for color, cols in color_groups.items():
                        print(f'Grouping for color {color}: {cols}')
                        if color is None or len(cols) == 0:
                            continue
                        if len(cols) == 1:
                            merged[cols[0]] = concept_matrix_reset[cols[0]]
                        else:
                            group_name = "/".join(cols)
                            print(f'Group name: {group_name}')
                            group_name_wrapped = wrap_label(group_name, width=self.get_wrap_width())
                            merged[group_name_wrapped] = concept_matrix_reset[cols].any(axis=1)
                    concept_matrix_reset = merged
                print(f'Final columns for upset: {concept_matrix_reset.columns}')
                upset_data = from_indicators(concept_matrix_reset, concept_matrix_reset.columns)
                print('UpSet data created')
                fig = plt.figure(figsize=(12, 8))
                upset = UpSet(upset_data, show_counts=True)
                axes = upset.plot(fig=fig)
                bar_ax = axes['intersections']
                matrix_ax = axes['matrix']
                bars = bar_ax.patches
                upset_index = upset_data.index
                wrap_width = self.get_wrap_width()
                # Restore color and label logic
                if self.use_colors.get():
                    yticks = matrix_ax.get_yticklabels()
                    wrapped_labels = []
                    for label in yticks:
                        concept = label.get_text()
                        wrapped = wrap_label(concept, width=wrap_width)
                        wrapped_labels.append(wrapped)
                    matrix_ax.set_yticklabels(wrapped_labels)
                    for label, concept in zip(matrix_ax.get_yticklabels(), [l.get_text().replace('\n', ' ') for l in yticks]):
                        if concept in self.color_mapping:
                            color = self.color_mapping[concept]
                        else:
                            first_concept = concept.split('/')[0]
                            color = self.color_mapping.get(first_concept, 'black')
                        label.set_color(color)
                        label.set_weight('bold')
                        label.set_fontsize(10)
                else:
                    yticks = matrix_ax.get_yticklabels()
                    wrapped_labels = [wrap_label(label.get_text(), width=wrap_width) for label in yticks]
                    matrix_ax.set_yticklabels(wrapped_labels)
                    for label in matrix_ax.get_yticklabels():
                        label.set_color('black')
                        label.set_weight('normal')
                        label.set_fontsize(10)
                # Add parameter value labels in red
                for i, (bar, intersection) in enumerate(zip(bars, upset_index)):
                    if i == 0:
                        continue
                    if bar.get_height() == 0:
                        continue
                    x = bar.get_x() + bar.get_width() / 2
                    intersection_idx = i - 1
                    actual_intersection = upset_index[intersection_idx]
                    mask = np.ones(len(concept_matrix), dtype=bool)
                    for col, present in zip(concept_matrix_reset.columns, actual_intersection):
                        original_cols = col.split('/') if '/' in col else [col]
                        if present:
                            for orig_col in original_cols:
                                if orig_col in concept_matrix.columns:
                                    mask &= concept_matrix[orig_col] == 1
                        else:
                            for orig_col in original_cols:
                                if orig_col in concept_matrix.columns:
                                    mask &= concept_matrix[orig_col] == 0
                    param_vals = concept_matrix.loc[mask, varying_param].unique()
                    def fmt(v):
                        try:
                            f = float(v)
                            return f"{f:.1f}"
                        except Exception:
                            return str(v)
                    label = ','.join(fmt(v) for v in param_vals) if len(param_vals) > 0 else ''
                    print(f"Drawing parameter label: '{label}' at x={x}")
                    y_label = len(concept_matrix_reset.columns) - 0.3
                    matrix_ax.text(x, y_label, label, ha='center', va='bottom', fontsize=10, color='red', rotation=0, clip_on=False, weight='bold')
                # Add title and subtitle
                other_params = [p for p in params if p != varying_param]
                fixed_vals = {param_labels[p]: subset[p].iloc[0] for p in other_params}
                fixed_str = ' | '.join(f"{p}={v}" for p, v in fixed_vals.items())
                plt.suptitle(f"UpSet Diagram: {param_labels[varying_param]} sweep\nOther Params: {fixed_str}", fontsize=14, y=0.98)
                plt.subplots_adjust(top=0.88, bottom=0.12)
                plt.tight_layout(rect=[0, 0.12, 1, 0.88])
                outpath = f"{outdir}/compare_{param_labels[varying_param]}_composed.png"
                try:
                    plt.savefig(outpath, dpi=150, bbox_inches='tight', pad_inches=0.5, format='png')
                except Exception as e:
                    print(f"Error saving PNG diagram: {e}")
                plt.close('all')
                output_files.append(outpath)
            self.root.after(0, self.update_results, output_files)
            self.generate_html_table(blocks, df, self.color_mapping, outdir, param_labels)
            self.generate_docx_and_csv(blocks, df, self.color_mapping, outdir, param_labels)
            self.generate_stats_files(blocks, df, self.color_mapping, outdir, param_labels)
            self.root.after(0, self.display_block_info)
        except Exception as e:
            print('Exception in process_data:')
            traceback.print_exc()
            self.root.after(0, self.show_error, str(e))
    
    def update_results(self, output_files):
        # Get current window width
        window_width = self.root.winfo_width()
        
        # Update zoomable canvas with new images
        self.zoomable_canvas.set_images(output_files, window_width)
        
        # Stop progress and update status
        self.progress.stop()
        self.process_btn.config(state='normal')
        self.status_label.config(text=f"Processing complete. Generated {len(output_files)} plots.")
        
        # Call fit_to_window after processing
        self.zoomable_canvas.fit_to_window()
    
    def show_error(self, error_msg):
        self.progress.stop()
        self.process_btn.config(state='normal')
        self.status_label.config(text=f"Error occurred: {error_msg}")
        messagebox.showerror("Processing Error", f"An error occurred during processing:\n\n{error_msg}")

    def generate_html_table(self, blocks, df, color_mapping, outdir, param_labels):
        html_lines = [
            "<style>",
            ".concept-box { display: inline-block; padding: 1px 3px; margin: 1px; border-radius: 2px; font-weight: bold; white-space: nowrap; font-size: 10px; color: #fff; }",
            ".parameter-box { font-weight: bold; color: #333; font-size: 11px; }",
            ".varying-parameter { background-color: #ffffcc; }",
            ".concepts-line { white-space: nowrap; overflow-x: auto; }",
            "table { width: 100%; border-collapse: collapse; margin-bottom: 20px; font-size: 12px; }",
            "th, td { border: 1px solid #ddd; padding: 6px; text-align: left; vertical-align: top; }",
            "th { background-color: #f2f2f2; font-size: 11px; }",
            "</style>",
            "<table>",
            "<thead><tr>"
            "<th class='varying-parameter'>Temperature</th>"
            "<th class='varying-parameter'>Top P</th>"
            "<th class='varying-parameter'>Top K</th>"
            "<th class='varying-parameter'>BM25 Weight</th>"
            "<th>Extracted Concepts</th>"
            "</tr></thead><tbody>"
        ]
        for start, end, varying_param in blocks:
            subset = df.iloc[start:end].copy()
            for idx, row in subset.iterrows():
                html_lines.append("<tr>")
                html_lines.append(f"<td class='varying-parameter'><span class='parameter-box'>{row.get('Temperature','')}</span></td>")
                html_lines.append(f"<td class='varying-parameter'><span class='parameter-box'>{row.get('Top-p','')}</span></td>")
                html_lines.append(f"<td class='varying-parameter'><span class='parameter-box'>{row.get('Top-k','')}</span></td>")
                html_lines.append(f"<td class='varying-parameter'><span class='parameter-box'>{row.get('BM25 Weight','')}</span></td>")
                html_lines.append("<td><div class='concepts-line'>")
                for concept in row['Concepts']:
                    color = color_mapping.get(concept, '#888')
                    html_lines.append(
                        f"<span class='concept-box' style='background-color: {color};'>{html.escape(concept)}</span>"
                    )
                html_lines.append("</div></td></tr>")
        html_lines.append("</tbody></table>")
        with open(os.path.join(outdir, "compare.htm"), "w", encoding="utf-8") as f:
            f.write("\n".join(html_lines))

    def get_wrap_width(self):
        try:
            return int(self.wrap_width_var.get())
        except Exception:
            return 75

    def merge_csv_files(self):
        import re
        import traceback
        file_paths = filedialog.askopenfilenames(
            title="Select CSV files to merge",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        if not file_paths:
            return
        # Store the folder name for DOCX title
        folder_name = os.path.basename(os.path.dirname(file_paths[0])) if file_paths else ""
        self.merged_folder_name = folder_name
        self.merge_label.config(text="Selected files:\n" + "\n".join([os.path.basename(fp) for fp in file_paths]), foreground="black")
        error_patterns = [
            r"error generating response:",
            r"api error occurred:",
            r"bad gateway",
            r"cloudflare",
            r"server disconnected without sending a response",
            r"getaddrinfo failed"
        ]
        error_found = False
        error_msgs = []
        dfs = []
        non_utf8_files = []
        for fp in file_paths:
            df = None
            last_exc = None
            used_encoding = None
            try:
                df = pd.read_csv(fp, encoding='utf-8')
                used_encoding = 'utf-8'
            except Exception as e:
                last_exc = e
                non_utf8_files.append(os.path.basename(fp))
                print(f"File {os.path.basename(fp)} could not be read as utf-8: {e}")
            if used_encoding and used_encoding != 'utf-8':
                print(f"Warning: File {os.path.basename(fp)} was read with encoding {used_encoding}, not utf-8. This may cause issues with special characters.")
            if df is None:
                error_found = True
                error_msgs.append(f"File: {os.path.basename(fp)}, Error reading file as utf-8: {last_exc}")
                continue
            for idx, row in df.iterrows():
                for col in df.columns:
                    val = str(row[col]).lower()
                    for pat in error_patterns:
                        if re.search(pat, val):
                            error_found = True
                            error_msgs.append(f"File: {os.path.basename(fp)}, Row: {idx+2}, Column: '{col}', Error: {row[col]}")
            dfs.append(df)
        if non_utf8_files:
            msg = "Merge aborted: The following files are not valid UTF-8 and may cause encoding issues. Please convert them to UTF-8 and try again:\n" + "\n".join(non_utf8_files)
            self.merge_label.config(text=msg, foreground="red")
            messagebox.showerror("Merge Error", msg)
            return
        if error_found:
            msg = "Merge aborted due to invalid data in the following locations:\n" + "\n".join(error_msgs)
            self.merge_label.config(text=msg, foreground="red")
            messagebox.showerror("Merge Error", msg)
            return
        try:
            merged_df = pd.concat(dfs, ignore_index=True)
            merged_path = os.path.join(os.getcwd(), "compare_input.csv")
            merged_df.to_csv(merged_path, index=False, encoding='utf-8')
            self.csv_file = merged_path
            self.file_label.config(text=os.path.basename(merged_path), foreground="black")
            self.merge_label.config(text=f"Merged:\n" + "\n".join([os.path.basename(fp) for fp in file_paths]) + f"\n(encoding: utf-8)", foreground="black")
            self.process_btn.config(state='normal')
        except Exception as e:
            import sys
            import traceback
            tb_str = traceback.format_exc()
            print(f"Failed to merge files: {e}\nTraceback:\n{tb_str}")
            error_detail = f"Failed to merge files: {e}\nFiles attempted: {', '.join([os.path.basename(fp) for fp in file_paths])}\nTraceback (see shell):\n{tb_str}"
            self.merge_label.config(text=error_detail, foreground="red")
            messagebox.showerror("Merge Error", error_detail)

    def generate_docx_and_csv(self, blocks, df, color_mapping, outdir, param_labels):
        try:
            import docx
            from docx.shared import RGBColor
            import csv
            doc = docx.Document()
            table = doc.add_table(rows=1, cols=5)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Temperature'
            hdr_cells[1].text = 'Top P'
            hdr_cells[2].text = 'Top K'
            hdr_cells[3].text = 'BM25 Weight'
            hdr_cells[4].text = 'Extracted Concepts'
            for start, end, varying_param in blocks:
                subset = df.iloc[start:end].copy()
                for idx, row in subset.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row.get('Temperature',''))
                    row_cells[1].text = str(row.get('Top-p',''))
                    row_cells[2].text = str(row.get('Top-k',''))
                    row_cells[3].text = str(row.get('BM25 Weight',''))
                    concept_strs = []
                    for concept in row['Concepts']:
                        color = color_mapping.get(concept, '#888')
                        concept_strs.append(f"{concept} [{color}]")
                    row_cells[4].text = ", ".join(concept_strs)
            docx_path = os.path.join(outdir, "compare.docx")
            doc.save(docx_path)
            csv_path = os.path.join(outdir, "compare.csv")
            with open(csv_path, "w", newline='', encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(['Temperature', 'Top P', 'Top K', 'BM25 Weight', 'Extracted Concepts'])
                for start, end, varying_param in blocks:
                    subset = df.iloc[start:end].copy()
                    for idx, row in subset.iterrows():
                        concept_strs = []
                        for concept in row['Concepts']:
                            color = color_mapping.get(concept, '#888')
                            concept_strs.append(f"{concept} [{color}]")
                        writer.writerow([
                            row.get('Temperature',''),
                            row.get('Top-p',''),
                            row.get('Top-k',''),
                            row.get('BM25 Weight',''),
                            ", ".join(concept_strs)
                        ])
        except Exception as e:
            print(f"Error generating DOCX/CSV: {e}")

    def generate_stats_files(self, blocks, df, color_mapping, outdir, param_labels):
        try:
            import pandas as pd
            import docx
            from docx.shared import RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_ORIENT
            from docx2pdf import convert as docx2pdf_convert
            import os
            txt_lines = []
            all_concepts = sorted(set(c for concepts in df['Concepts'] for c in concepts))
            # TXT
            txt_lines.append("Concepts Overview\n================\n")
            txt_lines.append(f"Total unique concepts: {len(all_concepts)}\n")
            txt_lines.append("\nAll Concepts:\n" + "\n".join(all_concepts) + "\n")
            txt_lines.append("\nConcepts per block:\n")
            for start, end, varying_param in blocks:
                subset = df.iloc[start:end].copy()
                block_concepts = sorted(set(c for concepts in subset['Concepts'] for c in concepts))
                txt_lines.append(f"{varying_param} ({start}-{end-1}): {len(block_concepts)} concepts\n" + ", ".join(block_concepts) + "\n")
            txt_lines.append("\nColor Groups:\n")
            color_groups = {}
            for concept, color in color_mapping.items():
                color_groups.setdefault(color, []).append(concept)
            for color, concepts in color_groups.items():
                txt_lines.append(f"Color {color}: {len(concepts)} concepts\n" + ", ".join(concepts) + "\n")
            txt_path = os.path.join(outdir, "compare_stats.txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write("".join(txt_lines))
            # DOCX
            doc = docx.Document()
            section = doc.sections[0]
            from docx.shared import Inches
            from docx.enum.section import WD_ORIENT
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap width and height for landscape
            new_width, new_height = section.page_height, section.page_width
            section.page_width = new_width
            section.page_height = new_height
            # Set narrow margins and header/footer
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.3)
            section.left_margin = Inches(0.3)
            section.right_margin = Inches(0.3)
            section.header_distance = Inches(0.3)
            section.footer_distance = Inches(0.3)
            # --- First page: folder name as title ---
            folder_name = getattr(self, 'merged_folder_name', '')
            if folder_name:
                doc.add_heading(folder_name, level=1)
            else:
                doc.add_heading("Merged Results", level=1)
            # Insert 4 images in a table, side by side, width matches table width
            image_dir = os.path.join(os.path.dirname(outdir), "compare_gui_output")
            image_files = [
                "compare_BM25_composed.png",
                "compare_Topk_composed.png",
                "compare_Topp_composed.png",
                "compare_Temp_composed.png"
            ]
            img_table = doc.add_table(rows=1, cols=4)
            img_table.autofit = False
            # Calculate available width and height for images (page width - margins)
            table_width = int(section.page_width - section.left_margin - section.right_margin)
            available_height = int(section.page_height - section.top_margin - section.bottom_margin)
            img_width = int(table_width // 4)
            from PIL import Image as PILImage
            for i, img_file in enumerate(image_files):
                img_path = os.path.join(image_dir, img_file)
                if os.path.exists(img_path):
                    cell = img_table.rows[0].cells[i]
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cell.paragraphs[0].add_run()
                    # Open image to get aspect ratio
                    with PILImage.open(img_path) as pil_img:
                        aspect = pil_img.width / pil_img.height
                        # Calculate height in inches to fill page height
                        height_in_inches = (available_height / 914400)
                        height_in_inches = max(0.5, height_in_inches - 0.1)
                        # Calculate width to maintain aspect ratio, but cap at img_width
                        width_in_inches = min(img_width / 914400, aspect * height_in_inches)
                        run.add_picture(img_path, width=docx.shared.Inches(width_in_inches))
            # --- Page break, then Color Groups section ---
            doc.add_heading("Color Groups", level=1)
            # Color Groups table: 4 columns, set widths proportionally to fill table_width
            color_table = doc.add_table(rows=1, cols=4)
            color_table.autofit = False
            col_props = [1.2, 0.4, 2.2, 3.2]
            total = sum(col_props)
            col_widths = [int(table_width * (w / total)) for w in col_props]
            for i, w in enumerate(col_widths):
                color_table.columns[i].width = w
            color_table.rows[0].cells[0].text = 'Group Label'
            color_table.rows[0].cells[1].text = 'Count'
            color_table.rows[0].cells[2].text = 'Concepts'
            color_table.rows[0].cells[3].text = 'Parameters'
            for color, concepts in color_groups.items():
                # Group label: common word in group (use the longest common substring or first word)
                group_label = ''
                if len(concepts) > 1:
                    from difflib import SequenceMatcher
                    def lcs(a, b):
                        match = SequenceMatcher(None, a, b).find_longest_match(0, len(a), 0, len(b))
                        return a[match.a: match.a + match.size]
                    lcs_str = concepts[0]
                    for c in concepts[1:]:
                        lcs_str = lcs(lcs_str, c)
                    group_label = lcs_str.strip() if lcs_str.strip() else concepts[0].split()[0]
                else:
                    group_label = concepts[0].split()[0]
                row_cells = color_table.add_row().cells
                row_cells[0].text = group_label
                row_cells[1].text = str(len(concepts))
                # Concepts (colored)
                para = row_cells[2].paragraphs[0]
                for i, concept in enumerate(concepts):
                    run = para.add_run(concept)
                    if color.startswith('#') and len(color) == 7:
                        r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                        run.font.color.rgb = RGBColor(r, g, b)
                    if i < len(concepts) - 1:
                        para.add_run(", ")
                # Parameters (expanded cell)
                param_strs = []
                for concept in concepts:
                    param_rows = df[df['Concepts'].apply(lambda lst: concept in lst)]
                    for _, row in param_rows.iterrows():
                        param_strs.append(f"(Temp: {row.get('Temperature','')}, Topp: {row.get('Top-p','')}, Topk: {row.get('Top-k','')}, BM25: {row.get('BM25 Weight','')})")
                row_cells[3].text = ", ".join(param_strs)
                # Color group label cell colored
                for para in row_cells[0].paragraphs:
                    for run in para.runs:
                        if color.startswith('#') and len(color) == 7:
                            r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                            run.font.color.rgb = RGBColor(r, g, b)
            # --- All Concepts Table ---
            doc.add_heading("All Concepts", level=1)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Index'
            hdr_cells[1].text = 'Concept'
            for idx, concept in enumerate(all_concepts):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx+1)
                para = row_cells[1].paragraphs[0]
                run = para.add_run(concept)
                color = color_mapping.get(concept, '#888')
                if color.startswith('#') and len(color) == 7:
                    r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                    run.font.color.rgb = RGBColor(r, g, b)
            # --- Concepts per Block Table ---
            doc.add_heading("Concepts per Block", level=1)
            block_table = doc.add_table(rows=1, cols=3)
            block_table.autofit = True
            block_table.rows[0].cells[0].text = 'Block'
            block_table.rows[0].cells[1].text = 'Concepts'
            block_table.rows[0].cells[2].text = 'Count'
            for start, end, varying_param in blocks:
                subset = df.iloc[start:end].copy()
                block_concepts = sorted(set(c for concepts in subset['Concepts'] for c in concepts))
                row_cells = block_table.add_row().cells
                row_cells[0].text = varying_param
                para = row_cells[1].paragraphs[0]
                for i, concept in enumerate(block_concepts):
                    run = para.add_run(concept)
                    color = color_mapping.get(concept, '#888')
                    if color.startswith('#') and len(color) == 7:
                        r, g, b = tuple(int(color[j:j+2], 16) for j in (1, 3, 5))
                        run.font.color.rgb = RGBColor(r, g, b)
                    if i < len(block_concepts) - 1:
                        para.add_run(", ")
                row_cells[2].text = str(len(block_concepts))
            # --- UpSet Histogram Section ---
            doc.add_heading("UpSet Diagram Histograms", level=1)
            # For each block, add a table with parameter, value, and histogram (bar heights/counts)
            for start, end, varying_param in blocks:
                subset = df.iloc[start:end].copy()
                doc.add_heading(f"{varying_param} Sweep", level=2)
                hist_table = doc.add_table(rows=1, cols=3)
                hist_table.autofit = True
                hist_table.rows[0].cells[0].text = 'Parameter'
                hist_table.rows[0].cells[1].text = 'Value'
                hist_table.rows[0].cells[2].text = 'Histogram (bar counts)'
                # For each unique value of the varying parameter, count the number of concepts present
                for val in sorted(subset[varying_param].unique()):
                    mask = subset[varying_param] == val
                    present_concepts = [c for concepts in subset[mask]['Concepts'] for c in concepts]
                    hist_row = hist_table.add_row().cells
                    hist_row[0].text = varying_param
                    hist_row[1].text = str(val)
                    hist_row[2].text = str(len(present_concepts))
            docx_path = os.path.join(outdir, "compare_stats.docx")
            doc.save(docx_path)
            # PDF: convert DOCX to PDF
            LIBREOFFICE_PATH = r"C:\\Program Files\\LibreOffice\\program\\soffice.exe"
            result = subprocess.run([
                LIBREOFFICE_PATH,
                "--headless",
                "--convert-to", "pdf",
                docx_path,
                "--outdir", outdir
            ], check=True, capture_output=True, text=True)
            pdf_path = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
            if not os.path.exists(pdf_path):
                raise FileNotFoundError(f"PDF not found at {pdf_path}")
            self.merge_label.config(text="Stats PDF generated successfully.", foreground="black")
        except Exception as e:
            print(f"Error generating stats TXT/PDF/DOCX: {e}")

    def on_group_by_subletters(self):
        if self.group_by_subletters.get():
            self.group_by_words.set(False)

    def on_group_by_words(self):
        if self.group_by_words.get():
            self.group_by_subletters.set(False)

    def on_color_criteria_toggle(self):
        if self.use_colors.get():
            self.group_by_same_color_cb.state(['!disabled'])
        else:
            self.group_by_same_color_cb.state(['disabled'])
            self.group_by_same_color.set(False)

def replace_bekker_with_greek(concept, bekker_map):
    for bekker, greek in bekker_map.items():
        if bekker in concept:
            # Replace the Bekker number with the Greek concept
            return concept.replace(bekker, greek)
    return concept

def convert_docx_to_pdf_libreoffice(docx_path, output_dir=None):
    if output_dir is None:
        output_dir = os.path.dirname(docx_path)
    try:
        result = subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf",
            docx_path,
            "--outdir", output_dir
        ], check=True, capture_output=True, text=True)
        print("LibreOffice output:", result.stdout)
        pdf_path = os.path.join(output_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF not found at {pdf_path}")
        return pdf_path
    except Exception as e:
        print("Error converting DOCX to PDF with LibreOffice:", e)
        return None

def main():
    root = tk.Tk()
    app = UpSetGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()